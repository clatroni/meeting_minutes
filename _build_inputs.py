"""One-shot generator for input/Rules.docx and input/Template.docx.

Run once (or whenever you want to reset them to the defaults):
    python _build_inputs.py

After that, edit Rules.docx and Template.docx directly in Word — run.py will
read your edits on every run.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

INPUT = Path(__file__).parent / "input"
INPUT.mkdir(exist_ok=True)

GREEN = RGBColor(0x86, 0xBC, 0x25)
DARK = RGBColor(0x00, 0x00, 0x00)
GRAY = RGBColor(0x53, 0x56, 0x5A)
LIGHT_GRAY = RGBColor(0x95, 0x95, 0x95)
RED = RGBColor(0xC6, 0x28, 0x28)
AMBER = RGBColor(0xE6, 0x51, 0x00)


def _band(doc, text, size=14):
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = GREEN


def _set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


# ==================================================================
# Rules.docx — read by the AI engine on every run
# ==================================================================
def build_rules():
    from datetime import datetime

    doc = Document()
    _set_default_font(doc)

    GREEN = RGBColor(0x86, 0xBC, 0x25)

    # ----- Cover header -----
    eyebrow = doc.add_paragraph()
    er = eyebrow.add_run("GUIDELINES")
    er.bold = True
    er.font.size = Pt(10)
    er.font.color.rgb = GREEN
    eyebrow.paragraph_format.space_after = Pt(2)

    title = doc.add_paragraph()
    tr = title.add_run("MoM Writing Rules")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = DARK
    title.paragraph_format.space_after = Pt(2)

    subtitle = doc.add_paragraph()
    sr = subtitle.add_run("How the AI cleans your transcript and writes the minutes")
    sr.font.size = Pt(13)
    sr.font.color.rgb = GRAY
    subtitle.paragraph_format.space_after = Pt(14)

    # Intro paragraph
    intro = doc.add_paragraph()
    ir = intro.add_run(
        "This document is the single source of truth for how meeting minutes are written. "
        "Every time you generate a MoM, the AI reads this file fresh — so editing here changes the next run. "
        "Step 1 governs how raw transcript noise is cleaned before extraction. "
        "Step 2 governs how the cleaned content is shaped into a professional MoM."
    )
    ir.font.size = Pt(11)
    intro.paragraph_format.space_after = Pt(20)

    cleanup = [
        ("Strip fillers",
         "Remove 'um', 'uh', 'erm', 'mmm', 'hmm', 'yeah' when standalone, 'you know', 'I mean', "
         "'like' as a filler, and 'so' as a sentence-opener filler. The cleaned transcript should "
         "read as professional prose."),
        ("Drop pleasantries",
         "Ignore openings ('good morning', 'hi all', 'thanks for joining'), goodbyes ('thanks', 'bye'), "
         "and side-chat ('can you hear me?', 'I was on mute', 'are you there?'). They carry no meeting content."),
        ("Drop noise markers",
         "Remove transcription artifacts in square brackets — [silence], [inaudible], [laughter], "
         "[crosstalk], [unclear] — along with any other bracketed annotations."),
        ("Consolidate same-speaker turns",
         "Merge consecutive utterances by the same speaker into one logical thought. The Teams export "
         "splits long contributions across multiple short rows; the MoM should not."),
        ("Repair false starts",
         "Treat 'we — we are waiting on…' as 'we are waiting on…'. Drop the false start, "
         "keep the completed phrase."),
        ("Handle Greek-English code-switching",
         "Greek consulting calls routinely mix Greek with English business terms (e.g., "
         "'Θα στείλω το deck μέχρι την Παρασκευή', 'escalate-αρε στον Ανδρέα'). Translate the "
         "Greek portions silently and produce the MoM entirely in English. Keep widely-used English "
         "business terms intact (escalate, dry run, sign-off, deck, sprint, KPI, sizing, capacity)."),
        ("Greek fillers and pleasantries to drop",
         "In Greek transcripts, also drop: ε, λοιπόν, δηλαδή, νομίζω, βασικά, έτσι, εντάξει, "
         "ωραία, καλά, ναι, εμ, μμμ, σωστά, βεβαίως, φυσικά, ευχαριστώ, καλημέρα, γεια σας."),
        ("Resolve Greek vocative case",
         "Greek vocative drops the final 's' on male names. 'Νίκο' → 'Νικόλας' (rendered in English "
         "as 'Nikolas'). 'Γιάννη' → 'Γιάννης' ('Yannis'). 'Κώστα' → 'Κώστας' ('Kostas'). "
         "'Δημήτρη' → 'Δημήτρης' ('Dimitris'). 'Παύλο' → 'Παύλος' ('Pavlos'). Use the canonical "
         "form from the participant list."),
        ("Resolve speaker shorthand",
         "If someone is addressed as 'Niko' or 'Yanni' (Latin transliteration), map them back to "
         "the canonical full name ('Nikos Vasileiou', 'Yannis Karagiannis') from the participant list."),
        ("Ignore audio and connectivity issues",
         "Skip 'sorry I'm late', 'VPN issues', 'someone at the door', 'cutting in and out' — "
         "and Greek equivalents ('με ακούτε;', 'σπάει η σύνδεση', 'συγγνώμη που άργησα'). "
         "These are operational noise, not meeting content."),
    ]
    rules = [
        ("Output language",
         "Default to English. If the meeting is conducted entirely in Greek and the audience is "
         "Greek-speaking, the MoM may be written in Greek — but mixed-language outputs are forbidden. "
         "Pick one language and keep it consistent throughout."),
        ("Greek output vocabulary",
         "When producing a Greek MoM, use this exact vocabulary so every deliverable matches the "
         "team's house style. Section headers (uppercase): "
         "ΠΡΑΚΤΙΚΑ ΣΥΝΑΝΤΗΣΗΣ (eyebrow), ΤΙ ΣΥΖΗΤΗΘΗΚΕ, ΑΠΟΦΑΣΕΙΣ, "
         "ΕΠΟΜΕΝΑ ΒΗΜΑΤΑ & ΕΝΕΡΓΕΙΕΣ, ΚΙΝΔΥΝΟΙ & ΘΕΜΑΤΑ, ΣΗΜΑΝΤΙΚΕΣ ΗΜΕΡΟΜΗΝΙΕΣ, "
         "ΑΝΟΙΧΤΑ ΖΗΤΗΜΑΤΑ. Header info labels (mixed case): Πελάτης, Συμμετέχοντες, Στόχος, Γλώσσα. "
         "Action-item table columns: Ενέργεια, Υπεύθυνος, Προθεσμία, Προτεραιότητα, Κατάσταση. "
         "Priority values: Υψηλή / Μεσαία / Χαμηλή. Status values: Δεν Ξεκίνησε / Σε Εξέλιξη / Ολοκληρώθηκε. "
         "Risk types: Κίνδυνος / Θέμα. Risks table columns: Τύπος, Περιγραφή, Επίπτωση, Υπεύθυνος. "
         "Timeline columns: Ορόσημο, Ημερομηνία. Use 'Εκκρεμεί' instead of 'TBD' when no due date is known. "
         "Footer: 'Αυστηρά Εμπιστευτικό' instead of 'Strictly Private & Confidential'."),
        ("Greek narrative phrasing",
         "For the Greek executive summary use passive consulting style: "
         "'Η συνάντηση επικεντρώθηκε στο/Η συνάντηση κάλυψε τα θέματα...' (focus); "
         "'Λήφθηκαν N ουσιώδεις αποφάσεις, με σημαντικότερη: ...' (decisions); "
         "'Καταγράφηκαν N ενέργειες, εκ των οποίων οι X είναι Υψηλής προτεραιότητας...' (actions); "
         "'Συνιστάται η προσοχή της διοίκησης σε...' / 'Επισημάνθηκαν N κίνδυνοι προς παρακολούθηση' (risks); "
         "'Το επόμενο σημαντικό ορόσημο είναι η ...' (timeline). "
         "Avoid first-person and avoid English business filler in Greek prose."),
        ("Detect transcript language",
         "Determine the dominant language of the transcript before writing. If the transcript is "
         "predominantly Greek (>70% Greek characters), label it 'Greek' in the meeting metadata. "
         "If predominantly English (<15% Greek), label 'English'. Otherwise label 'Mixed (Greek/English)'. "
         "Surface this in the meeting_info.language field so the reader knows the source language."),
        ("Detect client / company",
         "Identify the client or company the meeting concerns and populate meeting_info.client_name. "
         "Look first in the meeting title (e.g., 'Enerwave Portfolio Reporting' → client 'Enerwave'); "
         "if the title is generic, scan the transcript for the most-mentioned external company name. "
         "Use the canonical brand name (e.g., 'Enerwave', 'ERB', 'PPC', 'Eurobank'), not a casual short form. "
         "If no specific client is implied (internal Deloitte meeting), leave the field empty. "
         "The value is used to personalize the follow-up email subject and greeting."),
        ("Tone",
         "Executive consulting prose. Concise, professional, ready to send to a senior client. "
         "Past-tense narrative throughout. Reads like a formal deliverable, not a transcript."),
        ("Voice",
         "Use passive voice for narrative sections (executive summary, decisions, risks, key dates). "
         "Examples: 'Six decisions were taken', 'Action items have been recorded', 'A material risk "
         "has been identified', 'The next milestone is scheduled for May 22'. "
         "Action item descriptions remain in verb-led imperative form for clarity ('Update the spec', "
         "'Send the deck for sign-off') — this is the consulting standard for owner-attributed tasks. "
         "Decisions may be phrased actively when they restate a directive ('Switch to net revenue', "
         "'Treat June 10 as a hard internal deadline')."),
        ("Forbidden in output",
         "No 'umm', 'uh', no filler, no raw transcript quotes, no generic titles ('Topic A'). "
         "No action items without owners when an owner can be reasonably inferred. No questions in "
         "place of statements (an action is a directive, not a question)."),
        ("Materiality filter",
         "Before adding an item to the MoM, evaluate whether a senior client/stakeholder would "
         "actually want to read it. Drop the following: "
         "(a) Acknowledgments and confirmations like 'Will do', 'Got it', 'Sounds good', 'Sure'. "
         "(b) Logistical micro-tasks ('send a calendar hold', 'set up a meeting') unless the logistics "
         "themselves are material to delivery. "
         "(c) Agenda transitions ('let's move to the next item', 'shall we go to KPIs'). "
         "(d) Items that are too granular for an executive view (someone bringing a colleague to a follow-up "
         "is rarely material; a decision to split a KPI by category usually is). "
         "(e) Repetitions of items already captured in another section. "
         "When in doubt, prefer to leave items out — the MoM should read like a curated executive brief, not a "
         "complete to-do list. Better to surface five high-value action items than fifteen mixed ones."),
        ("Structure",
         "Header → What Was Discussed → Decisions → Next Steps & Action Items → Risks & Issues → "
         "Key Dates → Open Questions. Skip a section only if genuinely empty."),
        ("Header",
         "Meeting title, date in ISO format (YYYY-MM-DD), duration, and attendees inline using "
         "the 'Name (Organization)' format. Always use full names, never first-name only."),
        ("Executive summary",
         "Four to eight sentences of real prose. Lead with status, then key decisions, then risks, "
         "then next milestone. A senior reader should grasp the meeting in thirty seconds. "
         "Never produce a count recap ('5 participants exchanged 110 contributions')."),
        ("Discussion topics",
         "Three to six bullets. Each is a bold short title of three to six words followed by a "
         "one to three sentence summary. Cover what was discussed and what was concluded — not who "
         "said what. Merge restated points into a single topic. No duplicate titles. "
         "Skip pleasantry blocks (greetings, agenda transitions, time-availability talk). Never use "
         "the placeholder 'General Discussion' — every topic must have a substantive, descriptive title."),
        ("Voice and grammar",
         "Action items and decisions must be in third-person imperative form, never first person. "
         "Convert 'I'll send the spec' → 'Send the spec'; 'Θα στείλω follow-up' → 'Send follow-up'. "
         "Strip self-introductions ('Δημήτρη εδώ', 'Sorry, Maria here') — these belong to the speaker "
         "attribution, not the action description."),
        ("Resolve relative dates",
         "When the transcript references a relative day-of-week ('this Saturday', 'next Tuesday', "
         "'το Σάββατο', 'την Πέμπτη'), resolve it to the concrete calendar date relative to the "
         "meeting date. Render as 'Saturday 9 May 2026' (English) or 'Σάββατο 9 Μαΐου 2026' (Greek). "
         "Never leave a bare day-of-week reference in a final MoM."),
        ("Decisions",
         "A bulleted list, one entry per agreed decision. Bold the decision statement, add an "
         "optional italic rationale, and include the owner in parentheses. Capture each decision once. "
         "Strip conversational lead-ins ('Agreed.', 'Understood.', 'OK,'); keep only the substance."),
        ("Action items — phrasing",
         "Verb-led imperative form ('Escalate the procurement on Meteologica', not "
         "'Alex can you escalate so we don't slip the milestone?'). Specific and self-contained — "
         "the reader should understand the action without having read the transcript."),
        ("Action items — ownership",
         "Resolve every owner to the full name from the participant list ('Alex Petrou', not 'Alex'; "
         "'Nikos Vasileiou', not 'Niko'). When a speaker addresses someone ('Alex, can you …'), the "
         "addressee is the owner — not the speaker. When a speaker commits ('I will send the spec'), "
         "the speaker is the owner."),
        ("Action items — schema",
         "Every commitment becomes one row: Action (imperative) | Owner (full name) | "
         "Due (YYYY-MM-DD or a phrase like 'by Friday' / 'today' / 'EOD') | Priority | Status. Deduplicate — "
         "do not list the same commitment twice."),
        ("Action items — due date",
         "If no concrete deadline was discussed and none can be inferred from context, set the Due "
         "field to 'TBD'. Never leave it blank. Do not invent a date that was not mentioned."),
        ("Action items — status",
         "Status is one of: 'Not Started' (default for new commitments raised in this meeting), "
         "'In Progress' (work has begun), 'Completed' (task is done). Most rows in a fresh MoM are "
         "'Not Started' — only mark 'In Progress' or 'Completed' if the transcript explicitly indicates "
         "work has started or finished."),
        ("Priority logic",
         "High for items due today, EOD, this week, or escalated and blocking. Medium for items due "
         "in the next one to two weeks. Low for best-effort items with no firm date."),
        ("Risks and issues",
         "Type (Risk = forward-looking, Issue = current) | Description | Impact (High/Medium/Low) | Owner. "
         "Include only amber or red items — missing sign-off, blocker, slippage, capacity gap. "
         "Do not duplicate items already captured as action items; risks are observations, not commitments."),
        ("Key dates",
         "Concrete milestones with explicit dates, e.g. 'Sprint close · 2026-05-15' or "
         "'Board demo · 2026-06-18'. Skip if no concrete dates were discussed."),
        ("Open questions",
         "Genuinely unresolved items only — material questions where the team did not arrive at an "
         "answer. Skip rhetorical questions, pleasantries ('can you hear me?'), and already-answered items."),
    ]

    _write_step(doc, "Step 1", "Clean up the transcript before extracting", cleanup, GREEN)
    _write_step(doc, "Step 2", "Write the MoM following these rules", rules, GREEN)

    # Footer block
    doc.add_paragraph()
    div = doc.add_paragraph()
    dr = div.add_run("_" * 90)
    dr.font.size = Pt(6)
    dr.font.color.rgb = LIGHT_GRAY

    foot = doc.add_paragraph()
    fr = foot.add_run(
        f"Last regenerated {datetime.now().strftime('%Y-%m-%d')}  ·  "
        "Edit this file in Word and save — the next AI generation will pick up your changes."
    )
    fr.font.size = Pt(9)
    fr.font.color.rgb = LIGHT_GRAY
    fr.italic = True

    out = INPUT / "Rules.docx"
    doc.save(out)
    print(f"Wrote {out}")


def _write_step(doc, step_label: str, step_title: str, items: list[tuple[str, str]], green) -> None:
    """Render a numbered-step section with a green band header and one-paragraph rules."""
    # Step band header
    doc.add_paragraph()  # breathing room above
    band = doc.add_paragraph()
    sr = band.add_run(f"{step_label.upper()}    {step_title}")
    sr.bold = True
    sr.font.size = Pt(13)
    sr.font.color.rgb = green
    band.paragraph_format.space_after = Pt(2)

    # Thin underline strip
    underline = doc.add_paragraph()
    ur = underline.add_run("─" * 60)
    ur.font.size = Pt(8)
    ur.font.color.rgb = green
    underline.paragraph_format.space_after = Pt(8)

    for i, (label, body) in enumerate(items, 1):
        # Number + label on the same line, slightly larger
        head = doc.add_paragraph()
        nr = head.add_run(f"{i:02d}    ")
        nr.bold = True
        nr.font.size = Pt(11)
        nr.font.color.rgb = green
        lr = head.add_run(label)
        lr.bold = True
        lr.font.size = Pt(11)
        lr.font.color.rgb = DARK
        head.paragraph_format.space_after = Pt(2)

        # Body paragraph — indented, slightly muted
        body_p = doc.add_paragraph()
        br = body_p.add_run(body)
        br.font.size = Pt(11)
        br.font.color.rgb = GRAY
        body_p.paragraph_format.left_indent = Cm(0.85)
        body_p.paragraph_format.space_after = Pt(10)


# Helper imports
from docx.shared import Cm


# ==================================================================
# Template.docx — visual reference for the final output
# ==================================================================
def build_template():
    doc = Document()
    _set_default_font(doc)

    p = doc.add_paragraph()
    r = p.add_run("MoM Visual Template")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = DARK

    p = doc.add_paragraph()
    r = p.add_run(
        "This is a visual mockup of the final MoM. The renderer copies the styles and structure shown here. "
        "Edit fonts/colors in Word to change the look of all future MoMs."
    )
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY

    # Header mockup
    p = doc.add_paragraph()
    r = p.add_run("[Meeting Title]")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = DARK
    p = doc.add_paragraph()
    r = p.add_run("[Date]   ·   [Duration]")
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY
    p = doc.add_paragraph()
    lr = p.add_run("Attendees: ")
    lr.bold = True
    lr.font.size = Pt(11)
    p.add_run("[Name (Org), Name (Org), …]").font.size = Pt(11)

    _band(doc, "What Was Discussed")
    doc.add_paragraph(
        "[4–8 sentence executive summary. Lead with status, then key decisions, then risks. "
        "The reader should grasp the meeting in 30 seconds.]"
    )
    p = doc.add_paragraph()
    r = p.add_run("• [Topic title]: ")
    r.bold = True
    p.add_run("[1–3 sentence summary of what was discussed in this topic.]")

    _band(doc, "Decisions")
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run("[Decision statement]")
    r.bold = True
    r2 = p.add_run(" — [optional rationale]")
    r2.italic = True
    r3 = p.add_run("  ([Owner])")
    r3.font.color.rgb = GRAY

    _band(doc, "Next Steps & Action Items")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Action", "Owner", "Due", "Priority"]):
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for row, color in [
        (("[Action description]", "[Owner]", "[YYYY-MM-DD]", "High"), RED),
        (("[Action description]", "[Owner]", "[by Friday]", "Medium"), AMBER),
    ]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            if i == 3:
                run.bold = True
                run.font.color.rgb = color

    _band(doc, "Risks & Issues")
    p = doc.add_paragraph()
    r = p.add_run("(Section appears only if amber/red items were flagged.)")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY

    _band(doc, "Key Dates")
    p = doc.add_paragraph()
    r = p.add_run("(Section appears only if concrete dates were mentioned.)")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY

    _band(doc, "Open Questions")
    p = doc.add_paragraph()
    r = p.add_run("(Section appears only if there are genuinely unresolved items.)")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Strictly Private & Confidential")
    r.font.size = Pt(9)
    r.font.color.rgb = LIGHT_GRAY

    out = INPUT / "Template.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    build_rules()
    build_template()
