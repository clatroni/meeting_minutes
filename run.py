#!/usr/bin/env python3
"""AI Meeting Minutes Generator — lean single-file version.

Inputs (in input/):
  - Rules.docx        — the writing rules (the AI reads these every run)
  - Template.docx     — the visual template (styles carry into the output)
  - Transcript.docx   — the meeting transcript (also accepts .txt / .pdf / .vtt)

Output (in output/):
  - <date>_<meeting>_MoM.docx  — the final professional MoM, ready to send

Provider: Anthropic Claude if ANTHROPIC_API_KEY is set, else a deterministic
rule-based fallback so the pipeline always runs.

Usage:
    python run.py
"""
from __future__ import annotations

import json
import logging
import os
import re
import sys
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# ============================================================
# Paths
# ============================================================
BASE = Path(__file__).resolve().parent
INPUT = BASE / "input"
OUTPUT = BASE / "output"
RULES_DOC = INPUT / "Rules.docx"
TEMPLATE_DOC = INPUT / "Template.docx"

GREEN = RGBColor(0x86, 0xBC, 0x25)
DARK = RGBColor(0x00, 0x00, 0x00)
GRAY = RGBColor(0x53, 0x56, 0x5A)
LIGHT_GRAY = RGBColor(0x95, 0x95, 0x95)
RED = RGBColor(0xC6, 0x28, 0x28)
AMBER = RGBColor(0xE6, 0x51, 0x00)

# ============================================================
# Logging
# ============================================================
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")  # type: ignore[attr-defined]
except Exception:
    pass
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-7s | %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("mom")


# ============================================================
# Transcript loading + cleaning
# ============================================================
_NAME_START = r"[A-ZΑ-ΩΆΈΉΊΌΎΏΪΫ]"
TEAMS_LINE = re.compile(rf"^(?P<name>{_NAME_START}[\w\.\-' ]{{1,80}}?)\s{{2,}}(?P<ts>\d{{1,2}}:\d{{2}}(?::\d{{2}})?)\s*$", re.UNICODE)
TEAMS_LINE_LOOSE = re.compile(rf"^(?P<name>{_NAME_START}[\w\.\-' ]{{1,80}}?)\s+(?P<ts>\d{{1,2}}:\d{{2}}(?::\d{{2}})?)\s*$", re.UNICODE)
VTT_TAG = re.compile(r"<v\s+([^>]+?)>(.*)", re.IGNORECASE)
VTT_TIMING = re.compile(r"^\d{1,2}:\d{2}:\d{2}\.\d{3}\s+-->")
NOISE = re.compile(r"\s*\[[^\]]+\]\s*")
FILLERS = {"um", "uh", "umm", "uhh", "erm"}


def load_transcript(path: Path) -> tuple[str, str]:
    """Return (raw_text, format_suffix)."""
    suffix = path.suffix.lower()
    if suffix == ".txt" or suffix == ".vtt":
        return path.read_text(encoding="utf-8", errors="ignore"), suffix
    if suffix == ".docx":
        d = Document(str(path))
        return "\n".join(p.text for p in d.paragraphs), suffix
    if suffix == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        return "\n".join((p.extract_text() or "") for p in reader.pages), suffix
    raise ValueError(f"Unsupported transcript format: {suffix}")


def clean_transcript(raw: str, suffix: str) -> dict:
    title_hint, date_hint = None, None
    utterances: list[dict] = []

    if suffix == ".vtt":
        current_ts = ""
        in_note = False
        for line in raw.splitlines():
            s = line.strip()
            if not s or s == "WEBVTT":
                in_note = False
                continue
            if s.startswith("NOTE"):
                in_note = True
                continue
            if in_note:
                # Extract title/date from a NOTE header block
                if s.lower().startswith("meeting:") and title_hint is None:
                    title_hint = s.split(":", 1)[1].strip()
                elif s.lower().startswith("date:") and date_hint is None:
                    date_hint = s.split(":", 1)[1].strip()
                continue
            if VTT_TIMING.match(s):
                current_ts = s.split(" ")[0][:8]
                continue
            m = VTT_TAG.search(s)
            if m:
                utterances.append({"speaker": m.group(1).strip(), "ts": current_ts, "text": m.group(2).strip()})
            elif utterances:
                utterances[-1]["text"] += " " + s
    else:
        # Teams docx/txt format
        lines = raw.splitlines()
        # Title heuristic: first non-blank that doesn't look like a speaker line
        for s in lines[:8]:
            s = s.strip()
            if not s:
                continue
            if not (TEAMS_LINE.match(s) or TEAMS_LINE_LOOSE.match(s)):
                if title_hint is None:
                    title_hint = s
                elif date_hint is None and re.search(
                    r"\b(20\d{2}|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|"
                    r"monday|tuesday|wednesday|thursday|friday|"
                    # Greek months and days
                    r"ιαν|φεβ|μαρ|απρ|μα[ιΐ]|ιουν|ιουλ|αυγ|σεπ|οκτ|νοε|δεκ|"
                    r"δευτέρα|τρίτη|τετάρτη|πέμπτη|παρασκευή|σάββατο|κυριακή)\b",
                    s, re.IGNORECASE | re.UNICODE):
                    date_hint = s
                continue
            break

        current = None
        for line in lines:
            s = line.strip()
            if not s:
                if current and current["text"].strip():
                    utterances.append(current)
                    current = None
                continue
            m = TEAMS_LINE.match(s) or TEAMS_LINE_LOOSE.match(s)
            if m:
                name = m.group("name").strip()
                # speaker name plausibility: must contain at least 2 letters and capitalized first letter
                if name and name[0].isupper() and len(name) <= 80:
                    if current and current["text"].strip():
                        utterances.append(current)
                    current = {"speaker": name, "ts": m.group("ts").strip(), "text": ""}
                    continue
            if current is not None:
                current["text"] = (current["text"] + " " + s).strip() if current["text"] else s
        if current and current["text"].strip():
            utterances.append(current)

    # English mid-sentence fillers / discourse markers to strip
    en_fillers = [
        r"um+", r"uh+", r"erm+", r"mmm+", r"hmm+", r"ehm",
        r"you know", r"I mean", r"basically", r"sort of", r"kind of",
        r"actually", r"really", r"literally", r"obviously", r"honestly",
        r"frankly", r"to be honest", r"in fact", r"in any case",
        r"I guess", r"I suppose", r"let me see", r"let me think",
        r"as a matter of fact", r"at the end of the day",
    ]
    # Greek equivalents
    gr_fillers = [
        r"δηλαδή", r"βασικά", r"νομίζω", r"λοιπόν", r"ξέρεις", r"ξέρετε",
        r"καταλαβαίνεις", r"πάντως", r"γενικά", r"σου λέω", r"εν τέλει",
        r"στην ουσία",
    ]
    en_pattern = re.compile(rf"\b(?:{'|'.join(en_fillers)})\b[\s,]*", re.IGNORECASE)
    gr_pattern = re.compile(rf"\b(?:{'|'.join(gr_fillers)})\b[\s,]*", re.IGNORECASE | re.UNICODE)
    # Trailing question-tag fillers ('right?', 'you know?', 'σωστά;')
    tag_pattern = re.compile(r",\s*(?:right|you know|ok|okay|σωστά)\?\s*", re.IGNORECASE | re.UNICODE)

    # Polish each utterance — preserves raw_text so the UI can diff
    polished = []
    for u in utterances:
        raw_for_display = u["text"].strip()
        text = NOISE.sub(" ", u["text"])

        # Strip discourse fillers (English + Greek)
        text = en_pattern.sub(" ", text)
        text = gr_pattern.sub(" ", text)
        # Strip trailing question-tag fillers
        text = tag_pattern.sub(" ", text)

        # Strip leading discourse openers ('OK so', 'Right so', 'Well so', 'So,')
        text = re.sub(r"^(?:OK|Okay|Right|Well|So|Now|Look|Listen|Anyway)[,\s]+(?=[a-zα-ω])",
                      "", text, flags=re.IGNORECASE | re.UNICODE)

        # Repair false starts: "we — we are" / "the — the team" → second occurrence wins
        text = re.sub(r"\b(\w+)\s*[—–\-]\s*\1\b", r"\1", text, flags=re.IGNORECASE | re.UNICODE)
        # Repair simple stutter: "I I will" → "I will"
        text = re.sub(r"\b(\w+)\s+\1\b", r"\1", text, flags=re.IGNORECASE | re.UNICODE)

        # Strip "I think" when used as filler before a comma
        text = re.sub(r"\bI think,\s*", "", text, flags=re.IGNORECASE)

        # Tidy punctuation and whitespace
        text = re.sub(r"\s+([,.!?;])", r"\1", text)
        text = re.sub(r",\s*,", ",", text)
        text = re.sub(r"\s+", " ", text).strip(" ,")

        # Strip leading + trailing standalone fillers
        toks = text.split()
        while toks and toks[0].lower().strip(",.!?") in FILLERS:
            toks.pop(0)
        while toks and toks[-1].lower().strip(",.!?") in FILLERS:
            toks.pop()
        text = re.sub(r"\s+", " ", " ".join(toks)).strip()

        if not text:
            continue
        speaker = re.sub(r"\s*\([^)]+\)\s*$", "", u["speaker"]).strip()
        polished.append({
            "speaker": speaker,
            "ts": u["ts"],
            "text": text,                  # cleaned (fillers + false starts stripped)
            "raw_text": raw_for_display,   # original verbatim
        })

    speakers = sorted({u["speaker"] for u in polished})
    return {
        "utterances": polished,
        "speakers": speakers,
        "title_hint": title_hint,
        "date_hint": date_hint,
    }


def transcript_to_text(cleaned: dict) -> str:
    return "\n".join(f"{u['speaker']} [{u['ts']}]: {u['text']}" for u in cleaned["utterances"])


KNOWN_CLIENTS = [
    # Add or remove clients here — the matcher is case-insensitive.
    # Used to detect the client when a real transcript is dropped in.
    # Order matters only when two names overlap (longer first).
    "Enerwave",
    "LIDL",
    "Fourlis",
    "ERB",
    "EMI",
    "PPC",
    "Eurobank",
    "Alpha Bank",
    "Piraeus Bank",
    "National Bank",
    "OTE",
    "Cosmote",
    "Vodafone",
    "Mytilineos",
    "Public Power Corporation",
]


def detect_client(title: str, transcript_text: str = "") -> str:
    """Best-effort detection of the client/company the meeting concerns.

    Strategy:
      1. If a known client name appears in the meeting title, take it.
      2. Otherwise, count mentions in the transcript text and return the most
         frequently mentioned known client (must be mentioned >=2 times).
      3. Otherwise return empty string.
    """
    title = (title or "")
    body = (transcript_text or "")

    # Pass 1 — title match (preferred)
    for client in KNOWN_CLIENTS:
        if re.search(rf"\b{re.escape(client)}\b", title, re.IGNORECASE):
            return client

    # Pass 2 — transcript frequency
    counts: dict[str, int] = {}
    for client in KNOWN_CLIENTS:
        n = len(re.findall(rf"\b{re.escape(client)}\b", body, re.IGNORECASE))
        if n >= 2:
            counts[client] = n
    if counts:
        return max(counts, key=counts.get)

    return ""


def detect_language(cleaned: dict) -> str:
    """Detect the dominant language of the transcript by counting Greek vs Latin letters.

    Returns one of: 'Greek', 'English', 'Mixed (Greek/English)'.
    """
    text = " ".join(u.get("text", "") for u in cleaned.get("utterances", []))
    if not text:
        return "Unknown"
    greek = sum(1 for c in text if "Ͱ" <= c <= "Ͽ" or "ἀ" <= c <= "῿")
    latin = sum(1 for c in text if c.isalpha() and c.isascii())
    total = greek + latin
    if total == 0:
        return "Unknown"
    g_pct = greek / total
    if g_pct >= 0.7:
        return "Greek"
    if g_pct <= 0.15:
        return "English"
    return "Mixed (Greek/English)"


# ============================================================
# Rules loader
# ============================================================
def load_rules() -> str:
    if not RULES_DOC.exists():
        return ""
    d = Document(str(RULES_DOC))
    return "\n".join(p.text for p in d.paragraphs if (p.text or "").strip())


# ============================================================
# AI Engine — Anthropic Claude (preferred) + rule-based fallback
# ============================================================
SCHEMA = {
    "type": "object",
    "required": ["meeting_info", "executive_summary", "discussion_topics", "action_items", "decisions_log"],
    "properties": {
        "meeting_info": {
            "type": "object",
            "required": ["title", "date", "participants"],
            "properties": {
                "project_name": {"type": "string"},
                "client_name": {"type": "string", "description": "The client/company the meeting concerns (e.g., 'Enerwave'). Used for the email subject and greeting."},
                "title": {"type": "string"},
                "date": {"type": "string"},
                "duration": {"type": "string"},
                "objective": {"type": "string"},
                "language": {"type": "string"},
                "participants": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "required": ["name"],
                        "properties": {
                            "name": {"type": "string"},
                            "role": {"type": "string"},
                            "organization": {"type": "string"},
                        },
                    },
                },
            },
        },
        "executive_summary": {"type": "string"},
        "discussion_topics": {
            "type": "array",
            "items": {
                "type": "object",
                "required": ["title", "summary"],
                "properties": {"title": {"type": "string"}, "summary": {"type": "string"}},
            },
        },
        "decisions_log": {
            "type": "array",
            "items": {
                "type": "object",
                "required": ["decision"],
                "properties": {
                    "decision": {"type": "string"},
                    "rationale": {"type": "string"},
                    "owner": {"type": "string"},
                },
            },
        },
        "action_items": {
            "type": "array",
            "items": {
                "type": "object",
                "required": ["action", "owner", "priority", "status"],
                "properties": {
                    "action": {"type": "string"},
                    "owner": {"type": "string"},
                    "due_date": {"type": "string", "description": "YYYY-MM-DD or phrase ('by Friday'). Use 'TBD' if no deadline was discussed."},
                    "priority": {"type": "string", "enum": ["High", "Medium", "Low"]},
                    "status": {"type": "string", "enum": ["Not Started", "In Progress", "Completed"]},
                },
            },
        },
        "risks_issues": {
            "type": "array",
            "items": {
                "type": "object",
                "required": ["type", "description", "impact"],
                "properties": {
                    "type": {"type": "string", "enum": ["Risk", "Issue"]},
                    "description": {"type": "string"},
                    "impact": {"type": "string", "enum": ["High", "Medium", "Low"]},
                    "owner": {"type": "string"},
                },
            },
        },
        "timeline": {
            "type": "array",
            "items": {
                "type": "object",
                "required": ["milestone", "date"],
                "properties": {"milestone": {"type": "string"}, "date": {"type": "string"}},
            },
        },
        "open_questions": {"type": "array", "items": {"type": "string"}},
    },
}

SYSTEM_PROMPT_BASE = """You are an expert Project Manager, PMO Lead, and Consulting Delivery Manager. You convert raw Microsoft Teams meeting transcripts into structured, executive-grade Meeting Minutes (MoMs) for senior stakeholders.

Your output MUST follow the rules below (read fresh from input/Rules.docx every run):

=== TEAM RULES ===
{rules}
=== END TEAM RULES ===
{tone_block}
Submit your output via the `submit_minutes` tool. Fill every relevant field. Use empty arrays only when truly nothing applies."""

TONE_PRESETS = {
    "executive": (
        "\n=== TONE ===\n"
        "Executive mode. Keep the executive summary to 4 sentences. Use 3-5 short topic bullets. "
        "Action descriptions should be verb-led, under 18 words.\n"
        "=== END TONE ===\n"
    ),
    "detailed": (
        "\n=== TONE ===\n"
        "Detailed mode. Executive summary 8-12 sentences with full context. Use 5-8 topic bullets, each "
        "with a 2-4 sentence summary. Capture nuance and rationale.\n"
        "=== END TONE ===\n"
    ),
    "casual": (
        "\n=== TONE ===\n"
        "Internal/casual mode. Slightly less formal, suitable for an internal team channel. Still "
        "professional — no slang, but friendlier phrasing.\n"
        "=== END TONE ===\n"
    ),
    "default": "",
}


def call_anthropic(rules: str, cleaned: dict, transcript_text: str,
                   *,
                   model: str | None = None,
                   temperature: float | None = None,
                   max_tokens: int | None = None,
                   tone: str | None = None) -> dict:
    import anthropic
    client = anthropic.Anthropic()
    model = model or os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-6")
    temperature = 0.2 if temperature is None else temperature
    max_tokens = max_tokens or 8000
    tone_block = TONE_PRESETS.get((tone or "default").lower(), "")

    user_msg = (
        f"Speakers detected: {', '.join(cleaned['speakers'])}\n"
        f"Title hint: {cleaned.get('title_hint') or '(none)'}\n"
        f"Date hint: {cleaned.get('date_hint') or '(none)'}\n\n"
        f"=== TRANSCRIPT ===\n{transcript_text}\n=== END TRANSCRIPT ==="
    )

    log.info(f"Calling Anthropic ({model}, temp={temperature}, tone={tone or 'default'})...")
    resp = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        temperature=temperature,
        system=SYSTEM_PROMPT_BASE.format(rules=rules, tone_block=tone_block),
        tools=[{"name": "submit_minutes", "description": "Submit the structured Meeting Minutes.", "input_schema": SCHEMA}],
        tool_choice={"type": "tool", "name": "submit_minutes"},
        messages=[{"role": "user", "content": user_msg}],
    )
    for block in resp.content:
        if getattr(block, "type", None) == "tool_use":
            return block.input  # type: ignore[return-value]
    raise RuntimeError("Anthropic did not return the expected tool_use block")


# ----- Rule-based fallback -----
ACTION_HINTS = re.compile(
    # English commitments + permission requests
    r"\b(I'?ll|I will|we'?ll|we will|please|escalate|raise|update the spec|share|"
    r"will send|send a hold|export|expose|set up|put together|drop a|handle it)\b"
    # Greek equivalents — Θα στείλω/κάνω/χειριστώ (I'll send/do/handle), μπορείς να (can you), παρακαλώ (please)
    r"|\bΘα\s+(?:στείλ|κάν|χειριστ|μοιραστ|μιλήσ|ενημερώσ|ετοιμάσ|βάλ)\w*"
    r"|\bμπορείς\s+να\b|\bπαρακαλώ\b|\bπαρακαλώ\s+να\b"
    r"|\bθα\s+στείλω\b|\bθα\s+φροντίσω\b|\bθα\s+το\s+κάνω\b"
    r"|\bστείλ(?:ε|ετε|ει|τε)\b|\bμοιράσου\b|\bενημέρωσε\b|\bετοίμασε\b",
    re.IGNORECASE | re.UNICODE,
)

DECISION_HINTS = re.compile(
    r"\b(agreed|let's (?!send)|approved|default to|treat .* as|hard internal deadline|switch to)\b"
    # Greek: συμφωνώ (I agree), συμφωνήσαμε (we agreed), αποφασίζουμε (we decide), εντάξει (OK)
    r"|\bσυμφων(?:ώ|ούμε|ήσαμε|ώντας)\b|\bαποφασίζ(?:ουμε|ω)\b"
    r"|\bπάμε\s+με\b|\bμένει\s+έτσι\b|\bτο\s+κλείσαμε\b",
    re.IGNORECASE | re.UNICODE,
)

RISK_HINTS = re.compile(
    r"\b(risk|blocker|blocked|delay|slip|amber|red|escalat\w+|tight|at risk|no sign-off|missing|gap|concern)\b"
    # Greek: κίνδυνος, μπλοκαρισμένο, καθυστέρηση, πρόβλημα, ανησυχία
    r"|\bκίνδυν(?:ος|ο|οι|ων)\b|\bμπλοκαρισμέν\w+|\bκαθυστέρηση\b|\bπρόβλημα\b"
    r"|\bανησυχ(?:ώ|ία|ίες)\b|\bρίσκο\b|\bεμπόδιο\b",
    re.IGNORECASE | re.UNICODE,
)

QUESTION = re.compile(r"[?;]\s*$")  # Greek uses ';' for question mark

LEADING_FILLERS = re.compile(
    # English fillers
    r"^\s*(?:yeah|yes|ok|okay|mmm+|hmm+|right|sure|so|um+|uh+|well|look|listen|"
    r"alright|fair|fair fair|noted|got it|good|great|perfect|thanks|"
    r"agreed|understood|absolutely|exactly|correct|sounds good|will do|i agree|"
    r"that[\'’]?s right|that works"
    # Greek fillers — ε, ναι, εντάξει, ωραία, καλά, λοιπόν, δηλαδή, νομίζω, βασικά,
    # έτσι, κοίτα, άκου, εμ, μμμ, ωωω, γενικά, πάντως, πες, σωστά, ευχαριστώ, καλημέρα
    r"|ε|ναι|εντάξει|ωραία|καλά|λοιπόν|δηλαδή|νομίζω|βασικά|έτσι|κοίτα|άκου|"
    r"εμ|μμμ+|γενικά|πάντως|σωστά|ευχαριστώ|καλημέρα|γεια|γεια\s+σας|τέλεια|"
    r"συμφωνώ|μάλιστα|βεβαίως|σίγουρα|φυσικά|ας\s+πούμε)"
    r"\b[\s,.!?·\-—]*",
    re.IGNORECASE | re.UNICODE,
)


def _strip_filler(text: str) -> str:
    prev = None
    out = text.strip()
    while prev != out:
        prev = out
        out = LEADING_FILLERS.sub("", out).strip()
    if out:
        out = out[0].upper() + out[1:] if not out[0].isupper() else out
    return out


def _shorten(text: str, n: int) -> str:
    text = _strip_filler(text)
    return text if len(text) <= n else text[: n - 1].rsplit(" ", 1)[0] + "…"


# ============================================================
# Polishing helpers — turn verbatim transcript fragments into
# professional, client-ready prose.
# ============================================================
_VOCATIVE_MAP = {
    # Latin transliteration — Greek vocative drops the final 's'
    "Niko": "Nikos", "Yanni": "Yannis", "Kosta": "Kostas",
    "Vasili": "Vasilis", "Pavlo": "Pavlos", "Dimitri": "Dimitris",
    "Alex": "Alex", "Maria": "Maria", "Sofia": "Sofia",
    "Eleni": "Eleni", "Petros": "Petros",
    "Aresti": "Arestis", "Stelio": "Stelios", "Lefteri": "Lefteris",
    "Anesti": "Anestis", "Manoli": "Manolis", "Christo": "Christos",
    "Andrea": "Andreas", "Theodore": "Theodoros", "Yorgo": "Yorgos",
    # Greek script — vocative → nominative
    "Νίκο": "Νικόλας", "Γιάννη": "Γιάννης", "Κώστα": "Κώστας",
    "Δημήτρη": "Δημήτρης", "Παύλο": "Παύλος", "Πέτρο": "Πέτρος",
    "Βασίλη": "Βασίλης", "Ανδρέα": "Ανδρέας", "Στέλιο": "Στέλιος",
    "Λευτέρη": "Λευτέρης", "Μανώλη": "Μανώλης", "Άρη": "Άρης",
    "Νεκτάριε": "Νεκτάριος", "Σόφια": "Σοφία", "Ελένη": "Ελένη",
}

_ADDRESSEE_RE = re.compile(
    # English: "Alex, please send X" / "Alex can you send X"
    # Also handles a leading Capitalized name in either Latin or Greek script.
    r"^([A-ZΑ-ΩΆΈΉΊΌΎΏΪΫ][\wͰ-Ͽ]+)"
    r"(?:\s*[,:]\s+|\s+(?=please|can you|could you|would you|let[\'’]?s|shall|"
    r"παρακαλώ|μπορείς|μπορείτε|θέλεις\s+να|θα\s+πρέπει\s+να))",
    re.IGNORECASE | re.UNICODE,
)
_FIRST_PERSON_COMMIT_RE = re.compile(
    r"^\s*(?:I[\'’]?ll|I will|I[\'’]?d|I am going to|I[\'’]?m going to|"
    r"we[\'’]?ll|we will|we[\'’]?d|we are going to"
    # Greek: optionally preceded by 'Εγώ'/'Εμείς' personal pronoun
    r"|(?:εγώ\s+|εμείς\s+)?θα\s+(?:στείλ|κάν|φροντίσ|χειριστ|μοιραστ|μιλήσ|ενημερώσ|ετοιμάσ|βάλ|τρέξ)\w*"
    r"|(?:εγώ\s+|εμείς\s+)?θα\s+το\s+\w+"
    # Bare 'Εγώ' / 'Εμείς' at start before a verb
    r"|εγώ|εμείς)\s+",
    re.IGNORECASE | re.UNICODE,
)
_PERMISSION_RE = re.compile(
    r"^\s*(?:please|can you|could you|would you|let[\'’]?s|let us"
    # Greek: παρακαλώ (alone or with να), μπορείς να, ας
    r"|παρακαλώ(?:\s+να)?|μπορείς\s+να|μπορείτε\s+να|ας)\s+",
    re.IGNORECASE | re.UNICODE,
)
_TRAILING_QMARK = re.compile(r"[?;]\s*$")  # Greek uses ';' for question mark


def _resolve_owner(short_name: str, speakers: list[str]) -> str:
    """Map 'Alex' → 'Alex Petrou', 'Niko' → 'Nikos Vasileiou' using the speaker list.

    Returns "" if the candidate cannot be resolved to a real attendee — caller
    should treat this as 'no owner found' and fall back to the speaker.
    """
    if not short_name:
        return ""
    short = short_name.strip()
    candidate = _VOCATIVE_MAP.get(short, short)
    # Exact full name match
    for s in speakers:
        if s == candidate:
            return s
    # First-name match (case-insensitive)
    for s in speakers:
        if s.split()[0].lower() == candidate.lower():
            return s
    # Substring fallback (defensive)
    for s in speakers:
        if candidate.lower() in s.lower():
            return s
    return ""  # not a real attendee — caller decides what to do


_SELF_INTRO_RE = re.compile(
    # "Dimitris here." / "Δημήτρη εδώ" / "It's Christina here" / "Sorry, Maria here"
    r"^(?:sorry,?\s+)?(?:it[\'’]?s\s+)?[A-ZΑ-ΩΆΈΉΊΌΎΏΪΫ][\wά-ώΆ-Ώ]+\s+(?:here|εδώ)[.!,]?\s+",
    re.IGNORECASE | re.UNICODE,
)


def _extract_commitment_sentence(text: str) -> str:
    """When the utterance contains context + commitment, return the commitment sentence.

    Example: 'The request is already with IT. I'll send a follow-up today.'
             → 'I'll send a follow-up today.'

    Falls back to the original text if no commitment marker is found.
    """
    sentences = re.split(r"(?<=[.!?;])\s+", text.strip())
    sentences = [s for s in sentences if s.strip()]
    if len(sentences) <= 1:
        return text

    commit_re = re.compile(
        r"\b(I[\'’]?ll|I will|I[\'’]?d|we[\'’]?ll|we will|"
        # Greek future-commitment + imperative verb forms
        r"θα\s+(?:στείλ|κάν|φροντίσ|χειριστ|μοιραστ|μιλήσ|ενημερώσ|ετοιμάσ|βάλ|τρέξ)\w*|"
        r"(?:στείλ(?:ε|τε|ω)|ενημέρωσ(?:ε|τε|ω)|μοίρασ(?:ε|τε)|κάν(?:ε|τε)|"
        r"ετοίμασ(?:ε|τε)|βάλ(?:ε|τε)|μίλησ(?:ε|τε)|χειρίσου|"
        r"escalate-αρε|escalate))",
        re.IGNORECASE | re.UNICODE,
    )
    permission_re = re.compile(
        r"\b(please|can you|could you|παρακαλώ|μπορείς|μπορείτε|μπορούμε)",
        re.IGNORECASE | re.UNICODE,
    )
    chosen = [s for s in sentences if commit_re.search(s) or permission_re.search(s)]
    return " ".join(chosen) if chosen else text


def _polish_action(text: str, speaker: str, speakers: list[str]) -> tuple[str, str]:
    """Convert a verbatim utterance into an imperative action + resolved owner.

    Returns (clean_action_text, owner_full_name).
    """
    raw = text.strip()
    owner = speaker

    # -1) Strip self-introduction ('Δημήτρη εδώ', 'Sorry, Maria here')
    raw = _SELF_INTRO_RE.sub("", raw)

    # 0a) Pull out only the commitment sentence if there's surrounding context
    raw = _extract_commitment_sentence(raw)

    # 0b) Strip conversational lead-ins so the addressee regex can match
    raw = _strip_filler(raw)

    # 1) If addressed to someone else ("Alex, can you escalate...") → use them as owner
    addressee_match = _ADDRESSEE_RE.match(raw)
    if addressee_match:
        resolved = _resolve_owner(addressee_match.group(1), speakers)
        if resolved:  # only take the addressee if we can map them to a real attendee
            owner = resolved
            raw = raw[addressee_match.end() :]
        else:
            # bogus match (e.g. "Will send …") — keep speaker, don't strip
            owner = _resolve_owner(speaker.split()[0], speakers) if speaker else ""
            if not owner:
                owner = speaker
    elif re.search(r"\bI[\'’]?ll\b|\bI will\b|\bI[\'’]?d\b|\bI[\'’]?m going to\b|\bθα\s+(?:στείλ|κάν|φροντίσ|χειριστ|μοιραστ|μιλήσ|ενημερώσ|ετοιμάσ|βάλ)\w*", raw, re.IGNORECASE | re.UNICODE):
        owner = _resolve_owner(speaker.split()[0], speakers) if speaker else ""
        if not owner:
            owner = speaker

    # 2) Strip "How about" / "What about" leading questions
    raw = re.sub(r"^(?:how about|what about|how do you feel about)\s+", "", raw, flags=re.IGNORECASE)

    # 3) Strip first-person commitment prefix ("I'll send X" → "Send X" / "Θα στείλω" → "")
    raw = _FIRST_PERSON_COMMIT_RE.sub("", raw)

    # 4) Strip permission/imperative scaffolding ("please send X" → "Send X")
    raw = _PERMISSION_RE.sub("", raw)

    # 5) Drop trailing question marks (questions-as-commitments → statements)
    raw = _TRAILING_QMARK.sub(".", raw)

    # 6) Tidy up
    raw = re.sub(r"\s+", " ", raw).strip(" .,!?-—")
    if raw:
        raw = raw[0].upper() + raw[1:]
    return raw, owner


# Materiality filter — keep only items a senior client would care about
_TRIVIAL_ACTION_RE = re.compile(
    r"^(?:will do|got it|sure|noted|sounds good|of course|ok|okay|"
    r"absolutely|agreed|fair|fair fair|right|done|cool|alright|"
    r"thanks|thank you|will be in touch|i[\'’]?ll be there|"
    r"bye|goodbye|see you|talk later|"
    r"συμφωνώ|εντάξει|ok|σύμφωνη|σύμφωνος|θα τα πούμε|γεια)\.?$",
    re.IGNORECASE | re.UNICODE,
)
_TRIVIAL_QUESTION_RE = re.compile(
    r"^(?:two hours|three hours|one hour|half hour|next week|today)\??$",
    re.IGNORECASE,
)


def _is_material_action(text: str) -> bool:
    """Return True if the action item is worth surfacing to a senior client.

    Filters out trivial acknowledgments, single-word verbs with no object,
    and very short fragments that lack actionable substance.
    """
    if not text:
        return False
    stripped = text.strip().rstrip(".!?")
    if len(stripped.split()) < 4:
        return False
    if _TRIVIAL_ACTION_RE.match(stripped):
        return False
    if _TRIVIAL_QUESTION_RE.match(stripped):
        return False
    # Pure restatement with no commitment
    if re.match(r"^(yes|no|maybe|fine|good|great)\b", stripped, re.IGNORECASE) and len(stripped.split()) < 6:
        return False
    return True


def _is_material_decision(text: str) -> bool:
    if not text:
        return False
    stripped = text.strip().rstrip(".!?")
    if len(stripped.split()) < 4:
        return False
    if _TRIVIAL_ACTION_RE.match(stripped):
        return False
    return True


def _is_material_risk(text: str) -> bool:
    if not text:
        return False
    stripped = text.strip().rstrip(".!?")
    # Risks need real substance — at least 8 words to describe the issue
    if len(stripped.split()) < 8:
        return False
    return True


def _polish_decision(text: str) -> str:
    """Strip conversational lead-ins from decision text."""
    out = _strip_filler(text)
    # If it starts with "We'll X" / "I'll X" / "Let's X", strip the prefix
    out = _FIRST_PERSON_COMMIT_RE.sub("", out)
    out = _PERMISSION_RE.sub("", out)
    out = re.sub(r"\s+", " ", out).strip(" .,!?-—")
    if out:
        out = out[0].upper() + out[1:]
    return out


def _infer_owner(u: dict, speaker_set: set[str]) -> str:
    if re.search(r"\bI'?ll\b|\bI will\b", u["text"], re.IGNORECASE):
        return u["speaker"]
    m = re.search(r"^([A-Z][a-zA-Z]+)\s*[,:]\s*(?:please|can you)", u["text"])
    if m and m.group(1) in {s.split()[0] for s in speaker_set}:
        # match to full name
        for s in speaker_set:
            if s.split()[0] == m.group(1):
                return s
    m = re.search(r"\b([A-Z][a-zA-Z]+)\s+(?:please|can you)\b", u["text"])
    if m:
        for s in speaker_set:
            if s.split()[0] == m.group(1):
                return s
    return u["speaker"]


def _infer_priority(text: str) -> str:
    t = text.lower()
    if any(w in t for w in ["today", "asap", "blocker", "blocked", "escalate", "hard deadline", "hard internal", "eod"]):
        return "High"
    return "Medium"


def _extract_due(text: str) -> str:
    # ISO date "2026-05-22"
    m = re.search(r"\b(20\d{2}-\d{2}-\d{2})\b", text)
    if m:
        return m.group(1)
    # "by Monday", "by EOD", "by Friday", "by May 22" — but NOT "by more than"
    m = re.search(
        r"by\s+(EOD|EOW|today|tomorrow|next week|end of (?:next )?week|"
        r"Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday|"
        r"(?:May|June|July|April|March|February|January|August|September|October|November|December|"
        r"Feb|Jan|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{1,2})",
        text, re.IGNORECASE,
    )
    if m:
        return m.group(0).strip()
    # bare month-day
    m = re.search(r"\b(May|June|July|April|March|February|January|August|September|October|November|December|"
                  r"Feb|Jan|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{1,2}\b", text)
    return m.group(0) if m else ""


def _infer_impact(text: str) -> str:
    t = text.lower()
    if any(w in t for w in ["board", "exec", "cfo", "blocker", "amber", "red", "no sign-off", "slip"]):
        return "High"
    return "Medium"


_GR_MONTHS = {
    "ιανουαρίου": 1, "φεβρουαρίου": 2, "μαρτίου": 3, "απριλίου": 4, "μαΐου": 5, "μαιου": 5,
    "ιουνίου": 6, "ιουνιου": 6, "ιουλίου": 7, "ιουλιου": 7, "αυγούστου": 8, "αυγουστου": 8,
    "σεπτεμβρίου": 9, "σεπτεμβριου": 9, "οκτωβρίου": 10, "οκτωβριου": 10,
    "νοεμβρίου": 11, "νοεμβριου": 11, "δεκεμβρίου": 12, "δεκεμβριου": 12,
}


def _parse_date_hint(date_hint: str) -> str:
    """Best-effort 'YYYY-MM-DD' parse from a free-form date header.

    Handles English ('Wednesday, May 13, 2026, 10:00 AM') and Greek
    ('Πέμπτη, 14 Μαΐου 2026, 11:00 πμ'). Falls back to today's date when nothing parses.
    """
    if not date_hint:
        return datetime.now().strftime("%Y-%m-%d")

    # Try English: "May 13, 2026"
    m = re.search(r"([A-Za-z]{3,})\s+(\d{1,2}),?\s+(\d{4})", date_hint)
    if m:
        try:
            return datetime.strptime(f"{m.group(1)} {m.group(2)} {m.group(3)}", "%B %d %Y").strftime("%Y-%m-%d")
        except ValueError:
            try:
                return datetime.strptime(f"{m.group(1)} {m.group(2)} {m.group(3)}", "%b %d %Y").strftime("%Y-%m-%d")
            except ValueError:
                pass

    # Try Greek: "14 Μαΐου 2026" or "Πέμπτη, 14 Μαΐου 2026"
    m = re.search(r"(\d{1,2})\s+([Α-Ωα-ωΆ-Ώά-ώΐΰϊϋ]+)\s+(\d{4})", date_hint, re.UNICODE)
    if m:
        day = int(m.group(1))
        month_name = m.group(2).lower()
        year = int(m.group(3))
        mo = _GR_MONTHS.get(month_name)
        if mo:
            try:
                return datetime(year, mo, day).strftime("%Y-%m-%d")
            except ValueError:
                pass

    return datetime.now().strftime("%Y-%m-%d")


def _is_pleasantry_chunk(text: str) -> bool:
    """Return True when the chunk is mostly greetings, agenda transitions, or time-availability talk."""
    keywords = [
        "good morning", "good afternoon", "good evening",
        "thanks for joining", "thank you for joining", "thanks for making time",
        "shall we move", "shall we go to", "moving on to", "next item", "let's move",
        "I have an hour", "I have about", "minutes available today", "hour available today",
        "καλημέρα", "καλησπέρα", "ευχαριστώ που μπήκατε", "έχω περίπου",
        "λεπτά διαθέσιμα", "ώρα διαθέσιμη", "πάμε στο",
    ]
    low = text.lower()
    matches = sum(1 for k in keywords if k.lower() in low)
    return matches >= 2 and len(text.split()) < 80


def fallback_extract(cleaned: dict) -> dict:
    log.warning("Rule-based fallback (no ANTHROPIC_API_KEY). Output is functional but not AI-polished.")
    utts = cleaned["utterances"]
    speakers = set(cleaned["speakers"])
    title = cleaned.get("title_hint") or "Meeting"
    date_raw = cleaned.get("date_hint") or ""
    date = _parse_date_hint(date_raw)

    duration = ""
    if utts and utts[-1]["ts"]:
        parts = [int(p) for p in utts[-1]["ts"].split(":") if p.isdigit()]
        mins = parts[0] * 60 + parts[1] if len(parts) == 3 else parts[0] if parts else 0
        if mins:
            duration = f"{mins} min"

    speaker_list = list(speakers)

    # ----- Action items: imperative phrasing + resolved owners + materiality filter -----
    actions = []
    seen = set()
    action_source_texts = set()
    for u in utts:
        if ACTION_HINTS.search(u["text"]):
            polished, owner = _polish_action(u["text"], u["speaker"], speaker_list)
            if not _is_material_action(polished):
                continue
            polished = _shorten(polished, 200)
            key = polished[:80].lower()
            if key in seen:
                continue
            seen.add(key)
            action_source_texts.add(u["text"])
            actions.append({
                "action": polished,
                "owner": owner,
                "due_date": _extract_due(u["text"]) or "TBD",
                "priority": _infer_priority(u["text"]),
                "status": "Not Started",
            })

    # ----- Decisions: stripped conversational lead-ins + materiality filter -----
    decisions = []
    seen_d = set()
    for u in utts:
        if DECISION_HINTS.search(u["text"]) and len(u["text"]) > 25:
            polished = _polish_decision(u["text"])
            polished = _shorten(polished, 220)
            if not _is_material_decision(polished):
                continue
            key = polished[:80].lower()
            if key in seen_d:
                continue
            seen_d.add(key)
            owner_full = _resolve_owner(u["speaker"].split()[0], speaker_list) if u["speaker"] else ""
            decisions.append({"decision": polished, "rationale": "", "owner": owner_full or u["speaker"]})

    # ----- Risks: don't double-capture utterances already promoted to actions + materiality filter -----
    risks = []
    seen_r = set()
    for u in utts:
        if u["text"] in action_source_texts:
            continue
        if RISK_HINTS.search(u["text"]) and len(u["text"]) > 30:
            polished = _strip_filler(u["text"])
            polished = re.sub(r"\?\s*$", ".", polished).strip()
            polished = _shorten(polished, 220)
            if not _is_material_risk(polished):
                continue
            key = polished[:80].lower()
            if key in seen_r:
                continue
            seen_r.add(key)
            owner_full = _resolve_owner(u["speaker"].split()[0], speaker_list) if u["speaker"] else ""
            risks.append({
                "type": "Risk",
                "description": polished,
                "impact": _infer_impact(u["text"]),
                "owner": owner_full or u["speaker"],
            })

    # ----- Timeline -----
    timeline = []
    seen_t = set()
    for u in utts:
        for m in re.finditer(r"\b(?:by\s+)?(May|June|April|March|July|Feb|Jan|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+(\d{1,2})\b", u["text"]):
            ms = _shorten(_strip_filler(u["text"]), 180)
            key = ms[:80].lower()
            if key in seen_t:
                continue
            seen_t.add(key)
            timeline.append({"milestone": ms, "date": f"{m.group(1)} {m.group(2)}"})

    # ----- Open questions: filter out actions/risks already captured -----
    questions = []
    seen_q = set()
    for u in utts:
        if not QUESTION.search(u["text"]) or len(u["text"]) < 40:
            continue
        if u["text"] in action_source_texts:
            continue
        # skip pleasantries
        if re.search(r"^(can you hear|are you there|sorry|good morning|hi all)", u["text"], re.IGNORECASE):
            continue
        polished = _shorten(_strip_filler(u["text"]), 180)
        key = polished[:80].lower()
        if key in seen_q or len(polished) < 18:
            continue
        seen_q.add(key)
        questions.append(polished)
        if len(questions) >= 6:
            break

    # ----- Topics: chunk + dedup by title + materiality + better summaries -----
    topics_by_title: dict[str, list[str]] = {}
    chunk = max(1, len(utts) // 5)
    for i in range(0, len(utts), chunk):
        seg = utts[i:i + chunk]
        if len(seg) < 2:
            continue
        blob = " ".join(u["text"] for u in seg)
        if _is_pleasantry_chunk(blob):
            continue
        topic_title = _topic_title(blob)
        if topic_title == "General Discussion":
            # only keep this as a last-resort label; we'll filter later if better topics exist
            pass
        topics_by_title.setdefault(topic_title, []).append(blob)

    # If we have better-named topics, drop the generic 'General Discussion' bucket
    has_specific = any(t != "General Discussion" for t in topics_by_title)
    if has_specific:
        topics_by_title.pop("General Discussion", None)

    topics = []
    for topic_title, blobs in list(topics_by_title.items())[:5]:
        merged = " ".join(blobs)
        polished_blob = _strip_filler(merged)
        # First 2 substantive sentences in chronological order
        sentences = [s for s in re.split(r"(?<=[.!?])\s+", polished_blob) if len(s) > 30]
        summary = " ".join(sentences[:2]) if sentences else _shorten(polished_blob, 240)
        # Resolve relative day references ('this Saturday' → 'Saturday 9 May 2026')
        summary = _resolve_relative_dates(summary, date)
        topics.append({"title": topic_title, "summary": _shorten(summary, 280)})

    # ----- Executive summary: real prose from the structured items -----
    detected_language = detect_language(cleaned)
    exec_summary = _build_prose_summary(
        topics=topics,
        decisions=decisions,
        actions=actions,
        risks=risks,
        timeline=timeline,
        language=detected_language,
    )

    transcript_blob = " ".join(u.get("text", "") for u in utts)
    detected_client = detect_client(title, transcript_blob)

    return {
        "meeting_info": {
            "project_name": title.split("-")[0].strip() if "-" in title else title,
            "client_name": detected_client,
            "title": title,
            "date": date,
            "duration": duration,
            "objective": "Status sync and decision check-ins for the project.",
            "language": detected_language,
            "participants": [{"name": s, "role": "", "organization": ""} for s in cleaned["speakers"]],
        },
        "executive_summary": exec_summary,
        "discussion_topics": topics,
        "decisions_log": decisions,
        "action_items": actions,
        "risks_issues": risks,
        "timeline": timeline,
        "open_questions": questions,
    }


_DAYS_EN = {"monday": 0, "tuesday": 1, "wednesday": 2, "thursday": 3,
            "friday": 4, "saturday": 5, "sunday": 6}
_DAYS_GR = {"δευτέρα": 0, "τρίτη": 1, "τετάρτη": 2, "πέμπτη": 3,
            "παρασκευή": 4, "σάββατο": 5, "κυριακή": 6}


def _resolve_relative_dates(text: str, anchor_date: str) -> str:
    """Replace 'this Saturday' / 'το Σάββατο' / 'next Tuesday' with concrete dates.

    Anchor is the meeting date in YYYY-MM-DD. The next occurrence of the
    referenced day-of-week (after the anchor) is substituted.
    """
    from datetime import datetime as _dt, timedelta
    if not anchor_date:
        return text
    try:
        anchor = _dt.strptime(anchor_date, "%Y-%m-%d").date()
    except ValueError:
        return text

    def next_occurrence(target_dow: int, weeks_ahead: int = 0):
        delta = (target_dow - anchor.weekday()) % 7
        if delta == 0:
            delta = 7
        return anchor + timedelta(days=delta + weeks_ahead * 7)

    def en_format(d):
        # Cross-platform "Saturday 9 May 2026" — manual day formatting avoids %-d / %#d split
        return f"{d.strftime('%A')} {d.day} {d.strftime('%B')} {d.year}"

    out = text
    # English: "this Saturday", "next Saturday"
    for day_name, idx in _DAYS_EN.items():
        out = re.sub(
            rf"\bthis\s+{day_name}\b",
            lambda m, i=idx: en_format(next_occurrence(i)),
            out,
            flags=re.IGNORECASE,
        )
        out = re.sub(
            rf"\bnext\s+{day_name}\b",
            lambda m, i=idx: en_format(next_occurrence(i, weeks_ahead=1)),
            out,
            flags=re.IGNORECASE,
        )

    # Greek: "αυτό το Σάββατο" / "το Σάββατο" / "την Πέμπτη"
    gr_months = {1:"Ιανουαρίου",2:"Φεβρουαρίου",3:"Μαρτίου",4:"Απριλίου",5:"Μαΐου",
                 6:"Ιουνίου",7:"Ιουλίου",8:"Αυγούστου",9:"Σεπτεμβρίου",10:"Οκτωβρίου",
                 11:"Νοεμβρίου",12:"Δεκεμβρίου"}
    gr_days = {0:"Δευτέρα",1:"Τρίτη",2:"Τετάρτη",3:"Πέμπτη",4:"Παρασκευή",5:"Σάββατο",6:"Κυριακή"}

    def gr_format(d):
        return f"{gr_days[d.weekday()]} {d.day} {gr_months[d.month]} {d.year}"

    # Greek month names — used in lookahead to avoid replacing already-dated references
    gr_month_names = "|".join([
        "Ιανουαρίου", "Φεβρουαρίου", "Μαρτίου", "Απριλίου", "Μαΐου",
        "Ιουνίου", "Ιουλίου", "Αυγούστου", "Σεπτεμβρίου", "Οκτωβρίου",
        "Νοεμβρίου", "Δεκεμβρίου",
    ])
    for day_name, idx in _DAYS_GR.items():
        # 'αυτό το Σάββατο' / 'αυτή την Πέμπτη' → next occurrence
        out = re.sub(
            rf"\bαυτ[όή]\s+(?:το|την)\s+{day_name}\b",
            lambda m, i=idx: gr_format(next_occurrence(i)),
            out,
            flags=re.IGNORECASE | re.UNICODE,
        )
        # bare 'το/την Σάββατο' — but ONLY if NOT followed by an explicit date
        #   ('την Τετάρτη 25 Μαΐου' should not be re-resolved)
        out = re.sub(
            rf"\b(?:το|την)\s+{day_name}\b(?!\s+\d|\s+({gr_month_names}))",
            lambda m, i=idx: gr_format(next_occurrence(i)),
            out,
            flags=re.IGNORECASE | re.UNICODE,
        )

    return out


def _build_prose_summary(*, topics: list, decisions: list, actions: list,
                         risks: list, timeline: list, language: str = "") -> str:
    """Compose a passive-voice executive summary in the meeting's language."""
    if _lang_code(language) == "el":
        return _build_prose_summary_el(topics=topics, decisions=decisions, actions=actions,
                                       risks=risks, timeline=timeline)
    return _build_prose_summary_en(topics=topics, decisions=decisions, actions=actions,
                                   risks=risks, timeline=timeline)


def _build_prose_summary_en(*, topics, decisions, actions, risks, timeline) -> str:
    parts: list[str] = []
    if topics:
        unique = list(dict.fromkeys(t["title"] for t in topics if t.get("title")))
        if any(t != "General Discussion" for t in unique):
            unique = [t for t in unique if t != "General Discussion"]
        if len(unique) == 1:
            parts.append(f"The session focused on {unique[0]}.")
        elif len(unique) == 2:
            parts.append(f"The session covered {unique[0]} and {unique[1]}.")
        elif unique:
            front = ", ".join(unique[:-1])
            parts.append(f"The session covered {front}, and {unique[-1]}.")
    if decisions:
        n = len(decisions)
        material = decisions[0]["decision"].rstrip(".")
        if n == 1:
            parts.append(f"One material decision was taken: {material}.")
        else:
            parts.append(f"{n} material decisions were taken, the most significant being: {material}.")
    if actions:
        n = len(actions)
        high = sum(1 for a in actions if (a.get("priority") or "").lower() == "high")
        if high:
            parts.append(f"{n} action items have been recorded, of which {high} are rated High priority and warrant immediate attention.")
        else:
            parts.append(f"{n} action items have been recorded, each owner-attributed for follow-up.")
    if risks:
        high_r = [r for r in risks if (r.get("impact") or "").lower() == "high"]
        if high_r:
            parts.append(
                f"Executive attention is recommended on {len(high_r)} High-impact item(s); "
                f"most notably: {high_r[0]['description'].rstrip('.')[:120]}."
            )
        else:
            parts.append(f"{len(risks)} risk(s) have been flagged for monitoring; none are currently rated High.")
    if timeline:
        nxt = timeline[0]
        parts.append(f"The next concrete milestone is scheduled for {nxt.get('date', '')}.")
    if not parts:
        return "The session addressed routine project status; no material decisions, actions, or risks were captured."
    return " ".join(parts)


def _build_prose_summary_el(*, topics, decisions, actions, risks, timeline) -> str:
    """Greek passive-voice executive summary, matching consulting house style.

    Mirrors the example provided by the user — "Η συνάντηση επικεντρώθηκε...",
    "Συζητήθηκαν...", "Αποφασίστηκε...", "Υπήρξαν ανησυχίες...", "Το επόμενο βήμα είναι...".
    """
    parts: list[str] = []
    if topics:
        unique = list(dict.fromkeys(t["title"] for t in topics if t.get("title")))
        if any(t != "General Discussion" for t in unique):
            unique = [t for t in unique if t != "General Discussion"]
        if len(unique) == 1:
            parts.append(f"Η συνάντηση επικεντρώθηκε στο θέμα {unique[0]}.")
        elif len(unique) == 2:
            parts.append(f"Η συνάντηση κάλυψε τα θέματα {unique[0]} και {unique[1]}.")
        elif unique:
            front = ", ".join(unique[:-1])
            parts.append(f"Η συνάντηση κάλυψε τα θέματα {front}, καθώς και {unique[-1]}.")
    if decisions:
        n = len(decisions)
        material = decisions[0]["decision"].rstrip(".")
        if n == 1:
            parts.append(f"Λήφθηκε μία ουσιώδης απόφαση: {material}.")
        else:
            parts.append(f"Λήφθηκαν {n} ουσιώδεις αποφάσεις, με σημαντικότερη: {material}.")
    if actions:
        n = len(actions)
        high = sum(1 for a in actions if (a.get("priority") or "").lower() == "high")
        if high:
            parts.append(
                f"Καταγράφηκαν {n} ενέργειες, εκ των οποίων οι {high} είναι Υψηλής προτεραιότητας "
                f"και απαιτούν άμεση προσοχή."
            )
        else:
            parts.append(f"Καταγράφηκαν {n} ενέργειες, με σαφή ανάθεση υπευθύνου για κάθε μία.")
    if risks:
        high_r = [r for r in risks if (r.get("impact") or "").lower() == "high"]
        if high_r:
            parts.append(
                f"Συνιστάται η προσοχή της διοίκησης σε {len(high_r)} θέμα(τα) Υψηλής επίπτωσης· "
                f"κυρίως: {high_r[0]['description'].rstrip('.')[:120]}."
            )
        else:
            parts.append(
                f"Επισημάνθηκαν {len(risks)} κίνδυνος/οι προς παρακολούθηση· κανένας δεν αξιολογείται ως Υψηλής επίπτωσης."
            )
    if timeline:
        nxt = timeline[0]
        parts.append(f"Το επόμενο σημαντικό ορόσημο είναι η {nxt.get('date', '')}.")
    if not parts:
        return ("Η συνάντηση κάλυψε τη συνήθη ενημέρωση κατάστασης του έργου· "
                "δεν καταγράφηκαν ουσιώδεις αποφάσεις, ενέργειες ή κίνδυνοι.")
    return " ".join(parts)


def _topic_title(text: str) -> str:
    keywords = [
        # English
        ("data source", "Data Sources & Connectivity"),
        ("data quality", "Data Quality"),
        ("reconciliation", "Reconciliation"),
        ("kpi", "KPI Review & Design"),
        ("rollout", "Rollout & Deployment"),
        ("cutover", "Cutover & Go-Live"),
        ("rollback", "Cutover & Rollback"),
        ("demo", "Demo & Board Readiness"),
        ("license", "Licensing & Capacity"),
        ("capacity", "Licensing & Capacity"),
        ("uat", "UAT Progress"),
        ("rounding", "Financial Calculation Issues"),
        ("dry run", "Dry Run Planning"),
        ("steering", "Steering Committee Prep"),
        ("esg", "ESG Reporting"),
        ("availability", "Availability Metrics"),
        ("financial", "Financial KPI"),
        ("scope", "Scope & Approach"),
        ("governance", "Governance & Cadence"),
        ("risk", "Risks & Issues"),
        ("milestone", "Milestones & Dates"),
        ("integration", "Integration Status"),
        ("data model", "Data Model & Schema"),
        ("erd", "Data Model & Schema"),
        ("scada", "SCADA Integration"),
        # Greek
        ("ποιότητα δεδομέν", "Data Quality"),
        ("συμφωνί", "Reconciliation"),
        ("ανοιχτά issues", "Open Issues"),
        ("πελάτη", "Client Engagement"),
        ("κίνδυν", "Risks & Issues"),
        ("μπλοκαρισμέν", "Open Blockers"),
        ("πρόοδ", "Progress Update"),
        ("σχέδιο", "Plan & Approach"),
    ]
    low = text.lower()
    for k, t in keywords:
        if k in low:
            return t
    return "General Discussion"


# ============================================================
# Word renderer — clones Template.docx, fills in the content
# ============================================================
def render_mom(mom: dict, out_path: Path, *, sections: dict | None = None,
               template_path: Path | None = None) -> None:
    sections = sections or DEFAULT_SECTIONS
    tpl = template_path or TEMPLATE_DOC
    doc = None
    if tpl and tpl.exists():
        # Start from a copy of the template, then clear all body content
        # but inherit the styles (fonts, theme colors).
        try:
            doc = Document(str(tpl))
            body = doc.element.body
            for child in list(body):
                if child.tag.endswith("}sectPr"):
                    continue
                body.remove(child)
        except Exception as e:
            log.warning(f"Template {tpl.name} unavailable ({e}); rendering with built-in defaults.")
            doc = None
    if doc is None:
        doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    mi = mom.get("meeting_info", {})
    lang = mi.get("language", "") or ""

    # Document title strip (small green eyebrow)
    p = doc.add_paragraph()
    r = p.add_run(_L("eyebrow", lang))
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = GREEN

    # Meeting title
    p = doc.add_paragraph()
    r = p.add_run(mi.get("title", "Meeting"))
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(2)

    sub = "   ·   ".join([s for s in [mi.get("date", ""), mi.get("duration", "")] if s])
    if sub:
        p = doc.add_paragraph()
        r = p.add_run(sub)
        r.font.size = Pt(11)
        r.font.color.rgb = GRAY
        p.paragraph_format.space_after = Pt(8)

    # Header info table — compact two-column layout (Client / Attendees / Objective / Language)
    parts = mi.get("participants", [])
    if parts or mi.get("objective") or mi.get("language") or mi.get("client_name"):
        info_rows: list[tuple[str, str]] = []
        if mi.get("client_name"):
            info_rows.append((_L("client", lang), mi["client_name"]))
        if parts:
            info_rows.append((_L("attendees", lang), ", ".join(_fmt_part(x) for x in parts)))
        if mi.get("objective"):
            info_rows.append((_L("objective", lang), mi["objective"]))
        if mi.get("language"):
            info_rows.append((_L("language", lang), mi["language"]))

        info_table = doc.add_table(rows=len(info_rows), cols=2)
        info_table.autofit = False
        for ri, (label, value) in enumerate(info_rows):
            cell_label = info_table.rows[ri].cells[0]
            cell_value = info_table.rows[ri].cells[1]
            cell_label.text = ""
            cell_value.text = ""
            r = cell_label.paragraphs[0].add_run(label)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = GRAY
            r2 = cell_value.paragraphs[0].add_run(value)
            r2.font.size = Pt(11)
            _set_cell_width(cell_label, Cm(2.6))
        # Remove default borders
        _strip_table_borders(info_table)

    # What Was Discussed (executive summary + topics)
    show_summary = sections.get("executive_summary", True)
    show_topics = sections.get("discussion_topics", True)
    if (show_summary and mom.get("executive_summary")) or (show_topics and mom.get("discussion_topics")):
        _band(doc, _L("what_discussed", lang), lang)
        if show_summary and mom.get("executive_summary"):
            doc.add_paragraph(mom["executive_summary"])
        if show_topics:
            for t in mom.get("discussion_topics") or []:
                p = doc.add_paragraph()
                r = p.add_run(f"• {t.get('title', 'Topic')}: ")
                r.bold = True
                r.font.size = Pt(11)
                p.add_run(t.get("summary", "")).font.size = Pt(11)

    # Decisions
    if sections.get("decisions_log", True):
        decisions = mom.get("decisions_log") or []
        if decisions:
            _band(doc, _L("decisions", lang), lang)
            for d in decisions:
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(d.get("decision", ""))
                r.bold = True
                if d.get("rationale"):
                    r2 = p.add_run(f" — {d['rationale']}")
                    r2.italic = True
                if d.get("owner"):
                    r3 = p.add_run(f"  ({d['owner']})")
                    r3.font.color.rgb = GRAY

    # Next Steps & Action Items (always on — the headline)
    actions = mom.get("action_items") or []
    if actions:
        _band(doc, _L("next_steps", lang), lang)
        _table(doc,
               [_L("tbl_action", lang), _L("tbl_owner", lang), _L("tbl_due", lang),
                _L("tbl_priority", lang), _L("tbl_status", lang)],
               [[a.get("action", ""), a.get("owner", ""),
                 _due_or_tbd(a.get("due_date", ""), lang),
                 _translate_priority(a.get("priority", ""), lang),
                 _translate_status(a.get("status", ""), lang)]
                for a in actions],
               highlight_col=3, highlight_fn=_priority_color)

    # Risks
    if sections.get("risks_issues", True) and mom.get("risks_issues"):
        _band(doc, _L("risks", lang), lang)
        _table(doc,
               [_L("tbl_type", lang), _L("tbl_description", lang),
                _L("tbl_impact", lang), _L("tbl_owner", lang)],
               [[_translate_risk_type(r.get("type", ""), lang),
                 r.get("description", ""),
                 _translate_priority(r.get("impact", ""), lang),
                 r.get("owner", "")]
                for r in mom["risks_issues"]],
               highlight_col=2, highlight_fn=_priority_color)

    # Key Dates
    if sections.get("timeline", True) and mom.get("timeline"):
        _band(doc, _L("key_dates", lang), lang)
        _table(doc,
               [_L("tbl_milestone", lang), _L("tbl_date", lang)],
               [[t.get("milestone", ""), t.get("date", "")] for t in mom["timeline"]])

    # Open Questions
    if sections.get("open_questions", True) and mom.get("open_questions"):
        _band(doc, _L("open_questions", lang), lang)
        for q in mom["open_questions"]:
            doc.add_paragraph(q, style="List Bullet")

    # End-of-document classification line (above page footer)
    doc.add_paragraph()
    div = doc.add_paragraph()
    dr = div.add_run("_" * 90)
    dr.font.size = Pt(6)
    dr.font.color.rgb = LIGHT_GRAY

    p = doc.add_paragraph()
    r = p.add_run(_L("confidential", lang))
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY

    p = doc.add_paragraph()
    r = p.add_run(
        f"{_L('issued', lang)} {datetime.now().strftime('%Y-%m-%d')}  ·  "
        f"{_L('prepared_by', lang)}  ·  {_L('lab', lang)}"
    )
    r.font.size = Pt(9)
    r.font.color.rgb = LIGHT_GRAY

    # Page-footer with page number (proper Word section footer)
    _add_page_footer(doc, mi.get("title", "Meeting Minutes"))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def _add_page_footer(doc: Document, doc_title: str) -> None:
    """Add a real Word footer (left: doc title · right: page X of Y)."""
    try:
        section = doc.sections[0]
        footer = section.footer
        # Clear default content
        footer.is_linked_to_previous = False
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)

        para = footer.add_paragraph()
        # Left: document title
        left = para.add_run(doc_title[:80])
        left.font.size = Pt(8)
        left.font.color.rgb = LIGHT_GRAY

        # Right-tab to page number area
        para.add_run("\t\t")
        right = para.add_run("Page ")
        right.font.size = Pt(8)
        right.font.color.rgb = LIGHT_GRAY

        # Inject Word fields for { PAGE } and { NUMPAGES }
        for field_code in ("PAGE", "NUMPAGES"):
            fld_char_begin = OxmlElement("w:fldChar")
            fld_char_begin.set(qn("w:fldCharType"), "begin")
            instr_text = OxmlElement("w:instrText")
            instr_text.set(qn("xml:space"), "preserve")
            instr_text.text = f" {field_code} "
            fld_char_end = OxmlElement("w:fldChar")
            fld_char_end.set(qn("w:fldCharType"), "end")

            run_xml = right._r if field_code == "PAGE" else None
            if run_xml is None:
                # Insert " of " between PAGE and NUMPAGES
                of_run = para.add_run(" of ")
                of_run.font.size = Pt(8)
                of_run.font.color.rgb = LIGHT_GRAY
                run_xml = of_run._r
            run_xml.append(fld_char_begin)
            run_xml.append(instr_text)
            run_xml.append(fld_char_end)
    except Exception as e:
        log.warning(f"Could not add page footer: {e}")


def _band(doc, text, language=""):
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(_section_caps(text, language))
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = GREEN
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)


def _table(doc, headers, rows, highlight_col=None, highlight_fn=None):
    if not rows:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    # Header row — dark fill, white bold text
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, "1A1A1A")

    # Body rows — alternating subtle shading + priority/impact cell fills
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            # Subtle zebra striping on body rows
            if ri % 2 == 0:
                _shade_cell(cell, "FAFAFA")
            # Color-fill priority/impact cells
            if highlight_col == ci and highlight_fn:
                color, fill = highlight_fn(val)
                if color is not None:
                    run.font.color.rgb = color
                    run.bold = True
                if fill is not None:
                    _shade_cell(cell, fill)


def _shade_cell(cell, color_hex: str) -> None:
    """Apply background color to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _set_cell_width(cell, width) -> None:
    cell.width = width


def _strip_table_borders(table) -> None:
    """Hide borders on a layout-only table (used for the header info block)."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        return
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    # Remove existing borders if present
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def _fmt_part(p):
    if isinstance(p, str):
        return p
    name = p.get("name", "")
    role = p.get("role", "")
    org = p.get("organization", "")
    if role and org:
        return f"{name} ({role}, {org})"
    if role:
        return f"{name} ({role})"
    if org:
        return f"{name} ({org})"
    return name


def _priority_color(value):
    """Return (text_color, fill_hex) for a priority/impact cell. Bilingual."""
    v = (value or "").lower()
    if v in ("high", "υψηλή", "υψηλη"):
        return (RED, "FFEBEE")
    if v in ("medium", "μεσαία", "μεσαια"):
        return (AMBER, "FFF8E1")
    if v in ("low", "χαμηλή", "χαμηλη"):
        return (RGBColor(0x2E, 0x7D, 0x32), "E8F5E9")
    return (None, None)


def _slugify(text, max_len=60):
    text = re.sub(r"[^\w\s-]", "", text, flags=re.UNICODE).strip()
    text = re.sub(r"[\s_-]+", "_", text)
    return text[:max_len].strip("_") or "meeting"


# ============================================================
# Main
# ============================================================
def find_transcript() -> Path | None:
    """Pick the transcript file from input/ — anything that's not Rules.docx or Template.docx."""
    if not INPUT.exists():
        return None
    candidates = []
    for p in sorted(INPUT.iterdir()):
        if not p.is_file():
            continue
        if p.suffix.lower() not in {".docx", ".txt", ".pdf", ".vtt"}:
            continue
        if p.name in {"Rules.docx", "Template.docx"}:
            continue
        candidates.append(p)
    if not candidates:
        return None
    return candidates[0]


# ============================================================
# Bilingual labels — used by the renderer (Word + UI mockups).
# Data fields stay in English ('High', 'Risk', 'Not Started', 'TBD');
# the renderer translates when output language is Greek.
# ============================================================
LABELS = {
    "en": {
        "eyebrow":          "MEETING MINUTES",
        "what_discussed":   "What Was Discussed",
        "decisions":        "Decisions",
        "next_steps":       "Next Steps & Action Items",
        "risks":            "Risks & Issues",
        "key_dates":        "Key Dates",
        "open_questions":   "Open Questions",
        "client":           "Client",
        "attendees":        "Attendees",
        "objective":        "Objective",
        "language":         "Language",
        "tbl_action":       "Action",
        "tbl_owner":        "Owner",
        "tbl_due":          "Due",
        "tbl_priority":     "Priority",
        "tbl_status":       "Status",
        "tbl_type":         "Type",
        "tbl_description":  "Description",
        "tbl_impact":       "Impact",
        "tbl_milestone":    "Milestone",
        "tbl_date":         "Date",
        "tbd":              "TBD",
        "confidential":     "Strictly Private & Confidential",
        "issued":           "Issued",
        "prepared_by":      "Prepared by the AI Meeting Minutes Generator",
        "lab":              "Deloitte AI Prompting Lab",
        "high":             "High",
        "medium":           "Medium",
        "low":              "Low",
        "not_started":      "Not Started",
        "in_progress":      "In Progress",
        "completed":        "Completed",
        "risk":             "Risk",
        "issue":            "Issue",
    },
    "el": {
        "eyebrow":          "ΠΡΑΚΤΙΚΑ ΣΥΝΑΝΤΗΣΗΣ",
        "what_discussed":   "Τι Συζητήθηκε",
        "decisions":        "Αποφάσεις",
        "next_steps":       "Επόμενα Βήματα & Ενέργειες",
        "risks":            "Κίνδυνοι & Θέματα",
        "key_dates":        "Σημαντικές Ημερομηνίες",
        "open_questions":   "Ανοιχτά Ζητήματα",
        "client":           "Πελάτης",
        "attendees":        "Συμμετέχοντες",
        "objective":        "Στόχος",
        "language":         "Γλώσσα",
        "tbl_action":       "Ενέργεια",
        "tbl_owner":        "Υπεύθυνος",
        "tbl_due":          "Προθεσμία",
        "tbl_priority":     "Προτεραιότητα",
        "tbl_status":       "Κατάσταση",
        "tbl_type":         "Τύπος",
        "tbl_description":  "Περιγραφή",
        "tbl_impact":       "Επίπτωση",
        "tbl_milestone":    "Ορόσημο",
        "tbl_date":         "Ημερομηνία",
        "tbd":              "Εκκρεμεί",
        "confidential":     "Αυστηρά Εμπιστευτικό",
        "issued":           "Εκδόθηκε",
        "prepared_by":      "Δημιουργήθηκε από τον AI Meeting Minutes Generator",
        "lab":              "Deloitte AI Prompting Lab",
        "high":             "Υψηλή",
        "medium":           "Μεσαία",
        "low":              "Χαμηλή",
        "not_started":      "Δεν Ξεκίνησε",
        "in_progress":      "Σε Εξέλιξη",
        "completed":        "Ολοκληρώθηκε",
        "risk":             "Κίνδυνος",
        "issue":            "Θέμα",
    },
}


def _lang_code(language: str) -> str:
    """Return 'el' for Greek output, 'en' otherwise."""
    return "el" if "greek" in (language or "").lower() else "en"


def _L(key: str, language: str = "") -> str:
    """Return the localised label for the given language."""
    return LABELS[_lang_code(language)].get(key, LABELS["en"].get(key, key))


def _translate_priority(value: str, language: str) -> str:
    code = _lang_code(language)
    return {
        "high":   LABELS[code]["high"],
        "medium": LABELS[code]["medium"],
        "low":    LABELS[code]["low"],
    }.get((value or "").lower(), value or "")


def _translate_status(value: str, language: str) -> str:
    code = _lang_code(language)
    return {
        "not started": LABELS[code]["not_started"],
        "in progress": LABELS[code]["in_progress"],
        "completed":   LABELS[code]["completed"],
    }.get((value or "").lower(), value or LABELS[code]["not_started"])


def _translate_risk_type(value: str, language: str) -> str:
    code = _lang_code(language)
    return {
        "risk":  LABELS[code]["risk"],
        "issue": LABELS[code]["issue"],
    }.get((value or "").lower(), value or "")


# Greek tonos-stripping for section headers (Greek typography drops tonos in all-caps)
_GR_TONOS_MAP = str.maketrans({
    "ά": "α", "έ": "ε", "ή": "η", "ί": "ι", "ό": "ο", "ύ": "υ", "ώ": "ω",
    "ΐ": "ϊ", "ΰ": "ϋ",
    "Ά": "Α", "Έ": "Ε", "Ή": "Η", "Ί": "Ι", "Ό": "Ο", "Ύ": "Υ", "Ώ": "Ω",
})


def _section_caps(text: str, language: str = "") -> str:
    """Uppercase a section title using language-appropriate conventions."""
    if _lang_code(language) == "el":
        return text.translate(_GR_TONOS_MAP).upper()
    return text.upper()


def _due_or_tbd(due: str, language: str) -> str:
    """Render an empty / 'TBD' due-date as the language-appropriate placeholder."""
    s = (due or "").strip()
    if not s or s.upper() == "TBD":
        return _L("tbd", language)
    return s


DEFAULT_SECTIONS = {
    "executive_summary": True,
    "discussion_topics": True,
    "decisions_log": True,
    "action_items": True,  # always on (the headline)
    "risks_issues": True,
    "timeline": True,
    "open_questions": True,
}


def _maybe_review(rules: str, cleaned: dict, transcript_text: str, mom: dict,
                  *, model: str | None, max_tokens: int | None) -> dict:
    """Run the Michalis-style review pass when ENABLE_REVIEW_PASS is truthy. On any
    failure, return the original mom — review is quality polish, not load-bearing."""
    flag = os.environ.get("ENABLE_REVIEW_PASS", "1").strip().lower()
    if flag in ("0", "false", "no", "off", ""):
        return mom
    try:
        from modules.extractor_review import review_extraction
        reviewed = review_extraction(rules, cleaned, transcript_text, mom,
                                     schema=SCHEMA, model=model,
                                     max_tokens=max_tokens or 8000)
        log.info("Review pass: applied corrections.")
        return reviewed
    except Exception as e:
        log.warning(f"Review pass failed ({e}); keeping first extraction.")
        return mom


def _maybe_reconstruct(transcript_text: str, *, model: str | None,
                       max_tokens: int | None) -> str:
    """Run the Michalis-style LLM pre-cleaning pass when ENABLE_RECONSTRUCT is truthy.
    Returns the reconstructed transcript on success, or the original text on failure."""
    flag = os.environ.get("ENABLE_RECONSTRUCT", "0").strip().lower()
    if flag in ("0", "false", "no", "off", ""):
        return transcript_text
    try:
        from modules.reconstructor import reconstruct_transcript
        out = reconstruct_transcript(transcript_text, model=model,
                                     max_tokens=max_tokens or 8000)
        log.info(f"Reconstruct pass: {len(transcript_text)} → {len(out)} chars.")
        return out
    except Exception as e:
        log.warning(f"Reconstruct pass failed ({e}); using original cleaned text.")
        return transcript_text


def process_one(
    transcript_path: Path,
    output_dir: Path | None = None,
    *,
    provider: str | None = None,            # "auto" | "anthropic" | "rule_based" | None
    model: str | None = None,
    temperature: float | None = None,
    max_tokens: int | None = None,
    tone: str | None = None,                # "executive" | "detailed" | "casual" | "default"
    sections: dict | None = None,
    rules_path: Path | None = None,
    template_path: Path | None = None,
    filename_pattern: str | None = None,    # e.g. "{date}_{meeting}_MoM.docx"
) -> tuple[Path, dict]:
    """Process one transcript end-to-end with optional parameters.

    Used by both the CLI (run.py) and the Streamlit UI (ui.py).
    """
    rules_path = rules_path or RULES_DOC
    template_path = template_path or TEMPLATE_DOC
    if not rules_path.exists() or not template_path.exists():
        raise FileNotFoundError(
            f"Missing rules ({rules_path}) or template ({template_path}). "
            f"Run: python _build_inputs.py"
        )

    raw, suffix = load_transcript(transcript_path)
    cleaned = clean_transcript(raw, suffix)
    rules = _load_rules_from(rules_path)
    transcript_text = transcript_to_text(cleaned)

    use_provider = (provider or "auto").lower()
    if use_provider == "rule_based":
        mom = fallback_extract(cleaned)
    elif use_provider == "anthropic":
        transcript_text = _maybe_reconstruct(transcript_text, model=model,
                                             max_tokens=max_tokens)
        mom = call_anthropic(rules, cleaned, transcript_text,
                             model=model, temperature=temperature,
                             max_tokens=max_tokens, tone=tone)
        mom = _maybe_review(rules, cleaned, transcript_text, mom,
                            model=model, max_tokens=max_tokens)
    else:  # auto
        if os.environ.get("ANTHROPIC_API_KEY"):
            try:
                transcript_text = _maybe_reconstruct(transcript_text, model=model,
                                                     max_tokens=max_tokens)
                mom = call_anthropic(rules, cleaned, transcript_text,
                                     model=model, temperature=temperature,
                                     max_tokens=max_tokens, tone=tone)
                mom = _maybe_review(rules, cleaned, transcript_text, mom,
                                    model=model, max_tokens=max_tokens)
            except Exception as e:
                log.error(f"Anthropic call failed: {e}. Falling back to rule-based.")
                mom = fallback_extract(cleaned)
        else:
            mom = fallback_extract(cleaned)

    mi = mom.get("meeting_info", {})
    date = mi.get("date") or datetime.now().strftime("%Y-%m-%d")
    name = _slugify(mi.get("title", "Meeting"))
    out_dir = output_dir or OUTPUT
    pattern = filename_pattern or "{date}_{meeting}_MoM.docx"
    try:
        filename = pattern.format(date=date, meeting=name)
    except KeyError:
        filename = f"{date}_{name}_MoM.docx"
    out_path = out_dir / filename

    render_mom(mom, out_path, sections=sections or DEFAULT_SECTIONS, template_path=template_path)
    _save_transcript_sidecar(out_path, cleaned, mi)
    return out_path, mom


def _save_transcript_sidecar(mom_path: Path, cleaned: dict, meeting_info: dict) -> None:
    """Write a JSON sidecar next to the MoM: <stem>_transcript.json.

    Lets the UI show 'View transcript' for past minutes — so the user can
    audit what input produced each archived MoM.
    """
    sidecar = mom_path.with_name(mom_path.stem + "_transcript.json")
    payload = {
        "meeting_title": meeting_info.get("title", ""),
        "meeting_date": meeting_info.get("date", ""),
        "language": meeting_info.get("language", ""),
        "speakers": cleaned.get("speakers", []),
        "utterances": cleaned.get("utterances", []),
    }
    try:
        sidecar.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    except OSError as e:
        log.warning(f"Could not write transcript sidecar {sidecar.name}: {e}")


def _load_rules_from(path: Path) -> str:
    if not path.exists():
        return ""
    d = Document(str(path))
    return "\n".join(p.text for p in d.paragraphs if (p.text or "").strip())


def save_rules_text(text: str, path: Path | None = None) -> Path:
    """Persist a (possibly user-edited) rules text back to a Rules.docx file.

    Each non-empty line becomes one paragraph; lines starting with 'Step '
    render as section headers in green/bold. The format stays compatible with
    load_rules() so the AI engine can read the result on the next run.
    """
    target = path or RULES_DOC
    target.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    tr = title.add_run("MoM Writing Rules")
    tr.bold = True
    tr.font.size = Pt(20)

    note = doc.add_paragraph()
    nr = note.add_run("Edit this document — the AI reads it on every run.")
    nr.italic = True
    nr.font.size = Pt(10)
    nr.font.color.rgb = GRAY

    GREEN_RGB = RGBColor(0x86, 0xBC, 0x25)

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue
        # Section headers — lines starting with "Step "
        if line.lstrip().lower().startswith("step "):
            p = doc.add_paragraph()
            r = p.add_run(line.strip())
            r.bold = True
            r.font.size = Pt(13)
            r.font.color.rgb = GREEN_RGB
            continue
        # Bullet body lines — keep the leading '— ' / '• ' if present, otherwise plain
        p = doc.add_paragraph(style="List Bullet")
        # Bold-prefix support: "Label — value"
        if " — " in line and len(line.split(" — ", 1)[0]) <= 40:
            label, body = line.split(" — ", 1)
            r = p.add_run(f"{label} — ")
            r.bold = True
            r.font.size = Pt(11)
            p.add_run(body).font.size = Pt(11)
        else:
            p.add_run(line).font.size = Pt(11)

    doc.save(target)
    return target


def main() -> int:
    log.info("=" * 60)
    log.info("AI Meeting Minutes Generator")
    log.info("=" * 60)

    transcript_path = find_transcript()
    if transcript_path is None:
        log.error(f"No transcript found in {INPUT}/ (drop a .docx/.txt/.pdf/.vtt — anything besides Rules.docx and Template.docx)")
        return 1
    log.info(f"Transcript:  {transcript_path.name}")
    log.info(f"Rules:       {RULES_DOC.name}  ({len(load_rules())} chars of rules loaded)")
    log.info(f"Template:    {TEMPLATE_DOC.name}")

    try:
        out_path, mom = process_one(transcript_path)
    except FileNotFoundError as e:
        log.error(str(e))
        return 1

    log.info(f"Parsed and processed.")
    log.info("=" * 60)
    log.info(f"DONE -> {out_path}")
    log.info("=" * 60)
    return 0


if __name__ == "__main__":
    sys.exit(main())
