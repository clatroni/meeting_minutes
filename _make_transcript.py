"""Regenerate input/Transcript.docx (the messy realistic Enerwave call)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

OUT = Path(__file__).parent / "input" / "Transcript.docx"

LINES = [
    ("Christina Latroni", "0:00:03", "Good morning, good morning."),
    ("Christina Latroni", "0:00:05", "Uh, can everyone hear me OK? Niko?"),
    ("Nikos Vasileiou", "0:00:09", "Yeah, loud and clear Christina."),
    ("Sofia Mitropoulou", "0:00:11", "Hi all."),
    ("Christina Latroni", "0:00:13", "Great. Um, I think Alex is — Alex are you there?"),
    ("Alex Petrou", "0:00:18", "Sorry, I was on mute. Yes I'm here."),
    ("Christina Latroni", "0:00:23", "So Yannis is also on the call, he's been driving the data ingestion piece this sprint."),
    ("Yannis Karagiannis", "0:00:31", "Hi everyone."),
    ("Nikos Vasileiou", "0:00:33", "Καλημέρα Yanni. I have about an hour today."),
    ("Nikos Vasileiou", "0:00:40", "I'd like to spend most of it on the KPI review and then ten minutes on the rollout."),
    ("Christina Latroni", "0:00:48", "Yeah, that works for us."),
    ("Christina Latroni", "0:00:57", "Yannis can you walk us through where we are with the data sources?"),
    ("Yannis Karagiannis", "0:01:07", "Since last week we got three of the five sources connected."),
    ("Yannis Karagiannis", "0:01:13", "That's SCADA from the wind farms, the SAP PM extract, and the PPA contracts feed."),
    ("Yannis Karagiannis", "0:01:21", "The two still open are the weather API and the maintenance ticketing."),
    ("Alex Petrou", "0:01:29", "On the weather API, we're waiting on procurement. The Meteologica contract."),
    ("Alex Petrou", "0:01:36", "I pinged them yesterday. Realistic ETA is end of next week."),
    ("Christina Latroni", "0:01:43", "End of next week is tight for us. Our sprint closes May 15."),
    ("Christina Latroni", "0:01:51", "Alex can you escalate so we don't slip the milestone?"),
    ("Nikos Vasileiou", "0:01:57", "I'll handle it today."),
    ("Nikos Vasileiou", "0:02:01", "Worst case I authorize a short term data buy from a secondary provider."),
    ("Yannis Karagiannis", "0:02:12", "On the maintenance side — that one's bigger."),
    ("Yannis Karagiannis", "0:02:16", "ServiceNow has custom fields per plant and there's no schema documentation."),
    ("Yannis Karagiannis", "0:02:23", "I'd need somebody from your side for half a day to walk through it."),
    ("Sofia Mitropoulou", "0:02:30", "I can do it but not this week."),
    ("Sofia Mitropoulou", "0:02:35", "How about Monday afternoon? I'll bring Kostas from O&M."),
    ("Yannis Karagiannis", "0:02:43", "Monday afternoon is good. Two hours?"),
    ("Sofia Mitropoulou", "0:02:46", "Two hours yeah."),
    ("Yannis Karagiannis", "0:02:48", "I'll send a hold."),
    ("Christina Latroni", "0:02:52", "Good. Niko shall we move to the KPIs?"),
    ("Sofia Mitropoulou", "0:03:01", "I sent some feedback on the first draft."),
    ("Sofia Mitropoulou", "0:03:08", "The four quadrant view — availability, performance, financial, ESG — that's exactly how we want to brief the exec committee."),
    ("Sofia Mitropoulou", "0:03:18", "But I have three concerns."),
    ("Sofia Mitropoulou", "0:03:22", "First, the availability KPI."),
    ("Sofia Mitropoulou", "0:03:26", "Right now you're aggregating forced outages with planned maintenance and that's misleading."),
    ("Sofia Mitropoulou", "0:03:33", "Forced is a performance issue, planned isn't. They shouldn't be in the same number."),
    ("Yannis Karagiannis", "0:03:39", "Agreed. We'll add a toggle. Default to split."),
    ("Sofia Mitropoulou", "0:03:48", "Second, the financial KPI is using gross revenue."),
    ("Sofia Mitropoulou", "0:03:53", "We need to net out balancing costs and the curtailment compensations."),
    ("Sofia Mitropoulou", "0:04:00", "Otherwise the number is meaningless against the budget."),
    ("Nikos Vasileiou", "0:04:05", "Sofia is right. Exec committee will reject a gross figure on day one."),
    ("Christina Latroni", "0:04:11", "Understood. We'll switch to net."),
    ("Christina Latroni", "0:04:14", "Sofia can you share the formula your team uses for the netting?"),
    ("Sofia Mitropoulou", "0:04:23", "I'll send our internal one-pager today. CFO office approved it last year."),
    ("Sofia Mitropoulou", "0:04:34", "Third, ESG is just CO2 avoided right now. We need water consumption and waste recycling rate."),
    ("Sofia Mitropoulou", "0:04:42", "Both are in the new sustainability framework we adopted in March."),
    ("Yannis Karagiannis", "0:04:48", "Do you have those per asset, or just at corporate level?"),
    ("Sofia Mitropoulou", "0:04:54", "Per asset, monthly. They sit in HSE, not in SCADA. I'd need Alex to expose them."),
    ("Alex Petrou", "0:05:03", "HSE has a CSV export. There's no API."),
    ("Alex Petrou", "0:05:07", "I can set up a nightly drop to the Deloitte landing zone."),
    ("Yannis Karagiannis", "0:05:13", "Nightly CSV is fine for v1. We can move to API later."),
    ("Christina Latroni", "0:05:19", "So three KPI changes: split availability, net financial, expand ESG."),
    ("Christina Latroni", "0:05:27", "Yannis update the spec by Friday and send to Sofia for sign-off."),
    ("Yannis Karagiannis", "0:05:32", "Will do."),
    ("Nikos Vasileiou", "0:05:38", "The board has asked me to present the agent live in the June 18 portfolio review."),
    ("Nikos Vasileiou", "0:05:46", "So we need a working demo by June 10. With real data. From at least three plants."),
    ("Christina Latroni", "0:05:55", "June 10 is doable but tight."),
    ("Christina Latroni", "0:05:59", "It depends on the maintenance unblock. If Monday's session with Kostas is productive we can hit it."),
    ("Christina Latroni", "0:06:07", "If not, we're at risk."),
    ("Nikos Vasileiou", "0:06:13", "Let's treat June 10 as a hard internal deadline."),
    ("Nikos Vasileiou", "0:06:17", "Christina I'd like a written go/no-go from you on May 22."),
    ("Christina Latroni", "0:06:26", "May 22 go/no-go, agreed."),
    ("Christina Latroni", "0:06:29", "I'll send a one-pager with the recommendation and the underlying risks."),
    ("Sofia Mitropoulou", "0:06:37", "Quick question on rollout. Who actually uses this once it's live?"),
    ("Christina Latroni", "0:06:44", "Per the original scope, head office. But we built it mobile responsive in case you want to extend."),
    ("Sofia Mitropoulou", "0:06:53", "I'd like to extend. Asset managers spend 60% of their time on site."),
    ("Nikos Vasileiou", "0:07:05", "Sofia let's not expand scope mid-sprint. Phase two conversation."),
    ("Sofia Mitropoulou", "0:07:13", "Fair, OK."),
    ("Christina Latroni", "0:07:17", "Last item — the Power BI premium capacity."),
    ("Christina Latroni", "0:07:22", "The sandbox capacity isn't going to handle the volume once we're live."),
    ("Alex Petrou", "0:07:27", "I raised that internally on April 28. Procurement said two weeks. So Monday-ish."),
    ("Christina Latroni", "0:07:36", "If it slips by more than three days please flag it."),
    ("Alex Petrou", "0:07:43", "Will do."),
    ("Nikos Vasileiou", "0:08:13", "OK I think we covered everything."),
    ("Nikos Vasileiou", "0:08:16", "Christina send the recap by EOD?"),
    ("Christina Latroni", "0:08:19", "Yes you'll have minutes and actions in your inbox by 5."),
    ("Christina Latroni", "0:08:24", "Thanks everyone."),
    ("Sofia Mitropoulou", "0:08:26", "Thanks all."),
    ("Alex Petrou", "0:08:28", "Bye."),
    ("Yannis Karagiannis", "0:08:30", "Bye."),
    ("Nikos Vasileiou", "0:08:31", "Bye bye."),
]

GRAY = RGBColor(0x59, 0x59, 0x59)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

p = doc.add_paragraph()
r = p.add_run("Enerwave Portfolio Reporting - Weekly Sync")
r.bold = True
r.font.size = Pt(16)
p = doc.add_paragraph()
r = p.add_run("Wednesday, May 6, 2026, 11:00:00 AM (EEST)")
r.font.size = Pt(11)
r.font.color.rgb = GRAY
doc.add_paragraph().add_run("Transcript").bold = True
doc.add_paragraph()

for speaker, ts, text in LINES:
    p = doc.add_paragraph()
    r = p.add_run(speaker)
    r.bold = True
    r.font.size = Pt(11)
    p.add_run("   ")
    tr = p.add_run(ts)
    tr.font.size = Pt(10)
    tr.font.color.rgb = GRAY
    body = doc.add_paragraph(text)
    body.paragraph_format.space_after = Pt(4)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"Wrote {OUT}")
