"""Generate the three sample transcripts in input/transcript/.

Produces three professional, generic Teams .docx exports — no real-client
specifics, no overlap with confidential project content:

  01_Enerwave_ETL_Status_Sync.docx        — Data Engineering ETL project, 4 attendees, ~6 min
  02_Enerwave_Fabric_Backend_Review.docx  — Fabric reporting backend, 5 attendees, ~7 min
  03_LIDL_UAT_Exception_Review.docx       — PM + UAT exception review, 5 attendees, ~7 min (EL/EN mixed)

Run:
    python _make_samples.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).parent / "input" / "transcript"
ROOT.mkdir(parents=True, exist_ok=True)

GRAY = RGBColor(0x59, 0x59, 0x59)


def write_docx(path: Path, title: str, datetime_line: str, lines: list[tuple[str, str, str]]) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(16)
    p = doc.add_paragraph()
    r = p.add_run(datetime_line)
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY
    doc.add_paragraph().add_run("Transcript").bold = True
    doc.add_paragraph()

    for speaker, ts, text in lines:
        p = doc.add_paragraph()
        r = p.add_run(speaker)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run("   ")
        tr = p.add_run(ts)
        tr.font.size = Pt(10)
        tr.font.color.rgb = GRAY
        body = doc.add_paragraph(text)
        body.paragraph_format.space_after = Pt(4)

    doc.save(path)
    print(f"Wrote {path.name}  ({path.stat().st_size:,} bytes)")


# ============================================================
# Sample 01 — Enerwave ETL Pipeline Status Sync (English)
# Data Engineering project, generic source systems, transformation rules,
# data quality monitoring, go-live timeline.
# ============================================================
ETL_LINES = [
    ("Christina Latroni", "0:00:05", "Good morning everyone, thanks for joining the ETL status review."),
    ("Christina Latroni", "0:00:14", "We have four items on the agenda — source ingestion progress, the transformation layer, data quality checks, and the go-live timeline."),
    ("Eleni Antoniou", "0:00:28", "Morning Christina. I have a hard stop at the half-hour for an internal meeting."),
    ("Christina Latroni", "0:00:36", "Noted. Yannis, please walk us through the source ingestion progress."),
    ("Yannis Papadakis", "0:00:46", "Right, so we have completed connectivity to four of the six planned source systems."),
    ("Yannis Papadakis", "0:00:56", "The two remaining are the legacy operational database and the third-party billing API."),
    ("Yannis Papadakis", "0:01:08", "The operational database needs a service account from your infrastructure team. The billing API is blocked on a contract amendment."),
    ("Nikos Papadopoulos", "0:01:24", "I will request the service account today and have it provisioned by Friday."),
    ("Nikos Papadopoulos", "0:01:34", "On the billing API contract, procurement is reviewing this week. I expect sign-off by next Wednesday."),
    ("Christina Latroni", "0:01:46", "OK. If the contract slips beyond next Wednesday, we need a workaround — possibly a manual CSV drop in the interim."),
    ("Eleni Antoniou", "0:02:00", "Agreed. A manual CSV drop is acceptable for the first two weeks post go-live if needed."),
    ("Yannis Papadakis", "0:02:14", "On the transformation layer, we have implemented eighteen of the twenty-two business rules."),
    ("Yannis Papadakis", "0:02:26", "The four remaining are complex aggregations that need confirmation from your finance team."),
    ("Eleni Antoniou", "0:02:38", "I will arrange a session with the finance controller for Monday morning."),
    ("Christina Latroni", "0:02:48", "Monday works. Yannis, please prepare a one-pager with the four open rules and the proposed logic so we can validate quickly."),
    ("Yannis Papadakis", "0:03:02", "Will send the one-pager by end of day tomorrow."),
    ("Christina Latroni", "0:03:12", "On data quality — we discussed last week introducing a row-count and null-rate check at each stage."),
    ("Yannis Papadakis", "0:03:22", "Implemented. The checks run after each pipeline execution and write to a monitoring table."),
    ("Yannis Papadakis", "0:03:32", "First production-like run completed yesterday with no critical failures."),
    ("Christina Latroni", "0:03:42", "Excellent. Please share the monitoring dashboard with Eleni and Nikos this afternoon."),
    ("Yannis Papadakis", "0:03:52", "Will do."),
    ("Nikos Papadopoulos", "0:04:00", "One concern from our side — incremental load performance during the overnight window."),
    ("Nikos Papadopoulos", "0:04:10", "The full load takes around four hours. We need incremental loads completing within ninety minutes."),
    ("Yannis Papadakis", "0:04:22", "Current incremental load benchmarks at sixty minutes on production-equivalent volumes. Within target."),
    ("Christina Latroni", "0:04:34", "Good. Let us flag the four-hour full-load as a risk to monitor — if data volume doubles, we may exceed the window."),
    ("Christina Latroni", "0:04:50", "Last item — the go-live date. We are still tracking for the twenty-eighth of May."),
    ("Eleni Antoniou", "0:05:00", "From our side, business sign-off is expected by the twenty-second."),
    ("Christina Latroni", "0:05:10", "Then May the twenty-eighth remains the target. Yannis, please prepare a cutover checklist and share by next Tuesday."),
    ("Yannis Papadakis", "0:05:22", "Will prepare and share by Tuesday."),
    ("Christina Latroni", "0:05:32", "Anything else? OK, thanks everyone. Eleni, please drop when you need to."),
    ("Eleni Antoniou", "0:05:42", "Thanks Christina, talk Monday."),
]


# ============================================================
# Sample 02 — Enerwave Fabric Backend Review (English)
# Microsoft Fabric reporting project — capacity, OneLake, semantic model.
# ============================================================
FABRIC_LINES = [
    ("Christina Latroni", "0:00:05", "Good afternoon, thanks for joining the Fabric backend review for the reporting project."),
    ("Christina Latroni", "0:00:16", "Agenda — capacity sizing, OneLake structure, semantic model design, refresh strategy, and governance."),
    ("Stelios Vassiliou", "0:00:30", "Good afternoon Christina. Glad to walk through the design recommendations."),
    ("Christina Latroni", "0:00:38", "Dimitris, please start with capacity sizing."),
    ("Dimitris Kostas", "0:00:48", "Right, we recommend an F4 capacity for the initial rollout."),
    ("Dimitris Kostas", "0:00:58", "Volume estimates — around forty million rows per refresh, ten dashboards, fifty concurrent users. F4 gives us headroom for growth."),
    ("Stelios Vassiliou", "0:01:14", "What is the indicative monthly cost for F4?"),
    ("Dimitris Kostas", "0:01:22", "Around five thousand euro per month at list price. We can pause capacity outside business hours to reduce cost by approximately thirty percent."),
    ("Stelios Vassiliou", "0:01:36", "OK. Please add the pause-schedule cost saving to the business case."),
    ("Dimitris Kostas", "0:01:46", "Will update the business case by tomorrow."),
    ("Christina Latroni", "0:01:56", "Moving to OneLake structure — what is the proposed layout?"),
    ("Maria Georgiou", "0:02:06", "We propose the standard medallion architecture — Bronze for raw landed data, Silver for cleansed and conformed, Gold for the reporting marts."),
    ("Maria Georgiou", "0:02:20", "Each layer is a separate lakehouse for clear ownership and access control."),
    ("Stelios Vassiliou", "0:02:32", "Agreed on the medallion structure. Who owns each layer from your side?"),
    ("Maria Georgiou", "0:02:42", "Bronze and Silver are owned by the data engineering team. Gold is owned by the analytics team and exposed to business users."),
    ("Christina Latroni", "0:02:56", "Good. Let us document the ownership and access matrix and share with your IT team by Friday."),
    ("Maria Georgiou", "0:03:08", "Will prepare the matrix this week."),
    ("Christina Latroni", "0:03:18", "Semantic model — Dimitris, your recommendation?"),
    ("Dimitris Kostas", "0:03:26", "Direct Lake mode against the Gold lakehouse. No data duplication, minimal refresh management, sub-second query response for most dashboards."),
    ("Dimitris Kostas", "0:03:42", "Fallback to Import mode is configured for two complex DAX measures that hit Direct Lake limits."),
    ("Stelios Vassiliou", "0:03:56", "Are there any Direct Lake limitations we should flag to the business?"),
    ("Dimitris Kostas", "0:04:06", "Two main constraints — calculation groups behave differently, and certain DAX functions force a fallback to DirectQuery."),
    ("Dimitris Kostas", "0:04:20", "We have validated all twelve required dashboards against these limits, and ten run entirely in Direct Lake."),
    ("Christina Latroni", "0:04:32", "Good. Please document the two dashboards that use Import mode and the reason."),
    ("Dimitris Kostas", "0:04:44", "Will document and include in the architecture handover deck."),
    ("Christina Latroni", "0:04:54", "Refresh strategy?"),
    ("Maria Georgiou", "0:05:02", "OneLake tables refresh hourly during business hours through the data pipelines. The semantic model is auto-refreshed via Direct Lake."),
    ("Maria Georgiou", "0:05:16", "End-to-end latency from source system to dashboard is under fifteen minutes."),
    ("Stelios Vassiliou", "0:05:26", "Acceptable for our use case. Anything we need on the workspace setup?"),
    ("Maria Georgiou", "0:05:36", "Yes — we need three workspaces created. Development, Test, and Production."),
    ("Maria Georgiou", "0:05:46", "Plus a deployment pipeline configured to promote between them."),
    ("Stelios Vassiliou", "0:05:58", "Our platform team can provision those by next Wednesday."),
    ("Christina Latroni", "0:06:08", "Great. One risk to flag — Fabric capacity throttling under peak load."),
    ("Christina Latroni", "0:06:20", "If we hit one hundred concurrent users at month-end close, F4 may not be sufficient. We should plan a load test."),
    ("Dimitris Kostas", "0:06:34", "Will set up a load test during the UAT phase and report results."),
    ("Christina Latroni", "0:06:44", "Last item — governance. We need a workspace admin assigned and an approval flow for new reports."),
    ("Stelios Vassiliou", "0:06:56", "I will nominate the workspace admin by next Monday and propose an approval flow."),
    ("Christina Latroni", "0:07:08", "Excellent. Anything else?"),
    ("Stelios Vassiliou", "0:07:14", "Nothing from our side. Thanks for the clarity."),
    ("Christina Latroni", "0:07:20", "Thanks everyone. We will send the recap and the action list by end of day."),
]


# ============================================================
# Sample 03 — LIDL UAT Exception Review (Greek/English mixed)
# Project management confirmation + UAT exception update. Demonstrates
# bilingual conduct of a real engagement: Greek between Greek participants,
# English with German HQ stakeholder.
# ============================================================
LIDL_LINES = [
    ("Χριστίνα Λατρώνη", "0:00:05", "Καλημέρα σε όλους. Σήμερα έχουμε το UAT status review."),
    ("Στέφανος Βάλτερ", "0:00:14", "Καλημέρα Χριστίνα. Από το LIDL Headquarters συμμετέχουμε εγώ και η Άννα."),
    ("Anna Schmidt", "0:00:24", "Hello everyone. I will follow in English if that is acceptable."),
    ("Χριστίνα Λατρώνη", "0:00:32", "Of course Anna, we can mix languages as needed."),
    ("Χριστίνα Λατρώνη", "0:00:40", "Agenda σήμερα — UAT execution progress, defect summary, three exceptions που χρειάζονται απόφαση, και go-live readiness."),
    ("Χριστίνα Λατρώνη", "0:00:56", "Γιάννη, ξεκίνα με το status."),
    ("Γιάννης Παπαδάκης", "0:01:06", "Έχουμε εκτελέσει εκατόν είκοσι test cases από τα συνολικά εκατόν πενήντα."),
    ("Γιάννης Παπαδάκης", "0:01:18", "Από αυτά, ενενήντα δύο πέρασαν, είκοσι ένα απέτυχαν, και επτά είναι ακόμα σε εκτέλεση."),
    ("Στέφανος Βάλτερ", "0:01:30", "What is the severity breakdown of the failures?"),
    ("Γιάννης Παπαδάκης", "0:01:38", "Τέσσερις severity one, εννέα severity two, και οκτώ severity three."),
    ("Γιάννης Παπαδάκης", "0:01:50", "Τα τέσσερα severity one είναι όλα στο ίδιο module — την integration με το ERP backend."),
    ("Πέτρος Κωνσταντίνου", "0:02:02", "Confirm. The ERP integration team has acknowledged and a fix is expected by next Monday."),
    ("Χριστίνα Λατρώνη", "0:02:14", "OK, οπότε severity one defects θα κλείσουν Δευτέρα. Petros, παρακαλώ στείλε confirmation όταν deploy-εθεί το fix στο test environment."),
    ("Πέτρος Κωνσταντίνου", "0:02:28", "Θα στείλω confirmation Δευτέρα μεσημέρι, όταν ολοκληρωθεί το deployment."),
    ("Χριστίνα Λατρώνη", "0:02:40", "Καλά. Anna, ας συζητήσουμε τα τρία exceptions."),
    ("Anna Schmidt", "0:02:50", "Yes. The first exception concerns the multi-currency rounding rule."),
    ("Anna Schmidt", "0:03:00", "Our headquarters policy requires rounding at the invoice line, but the system rounds at the document total. The cumulative impact is roughly two cents per invoice."),
    ("Στέφανος Βάλτερ", "0:03:14", "We can accept this as a known issue if the cumulative impact stays under five hundred euro per month."),
    ("Χριστίνα Λατρώνη", "0:03:26", "Συμφωνώ. Καταγράφουμε ως exception, no change required. Γιάννη, βάλε το στο exception log."),
    ("Γιάννης Παπαδάκης", "0:03:38", "Θα το προσθέσω σήμερα στο exception log με σχετική τεκμηρίωση."),
    ("Anna Schmidt", "0:03:48", "Second exception — supplier master synchronization."),
    ("Anna Schmidt", "0:03:56", "There is a four-hour delay between supplier creation in the master system and availability in the front-end."),
    ("Πέτρος Κωνσταντίνου", "0:04:08", "This is a technical limitation of the source system. We can reduce to two hours with additional cost."),
    ("Στέφανος Βάλτερ", "0:04:20", "What is the additional cost?"),
    ("Πέτρος Κωνσταντίνου", "0:04:28", "Roughly fifteen hundred euro one-off for the configuration change, plus a small increase in compute cost."),
    ("Στέφανος Βάλτερ", "0:04:42", "Approved. Please proceed with the two-hour synchronization configuration."),
    ("Πέτρος Κωνσταντίνου", "0:04:52", "Θα κάνω submit το change request αυτή την εβδομάδα."),
    ("Anna Schmidt", "0:05:02", "Third exception — the discount approval workflow."),
    ("Anna Schmidt", "0:05:10", "The system requires manager approval for discounts above ten percent. Our policy allows discounts up to fifteen percent without approval for premium suppliers."),
    ("Στέφανος Βάλτερ", "0:05:24", "This needs a configuration change. Christina, can we add a rule based on supplier category?"),
    ("Χριστίνα Λατρώνη", "0:05:36", "Ναι, μπορούμε. Πέτρο, χρειαζόμαστε τη λίστα των premium suppliers με τα supplier IDs."),
    ("Πέτρος Κωνσταντίνου", "0:05:48", "Θα έχω τη λίστα μέχρι την Πέμπτη."),
    ("Χριστίνα Λατρώνη", "0:05:58", "Καλά. Maria, μπορείς να ετοιμάσεις το configuration change και να το έχουμε ready για test από επόμενη Δευτέρα;"),
    ("Μαρία Γεωργίου", "0:06:12", "Yes, I will prepare the configuration and have it ready for test next Monday."),
    ("Χριστίνα Λατρώνη", "0:06:24", "Τέλεια. Last item — go-live readiness."),
    ("Χριστίνα Λατρώνη", "0:06:32", "Με βάση τα open defects και τα exceptions, where do we stand?"),
    ("Γιάννης Παπαδάκης", "0:06:42", "Αν τα severity one κλείσουν Δευτέρα και τα exceptions έχουν resolution path, παραμένουμε στόχος για go-live στις δέκα Ιουνίου."),
    ("Στέφανος Βάλτερ", "0:06:56", "Acceptable from our side. We need a written go/no-go decision by the third of June."),
    ("Χριστίνα Λατρώνη", "0:07:08", "Συμφωνώ. Θα στείλω formal go/no-go assessment μέχρι τις τρεις Ιουνίου, με risk assessment και mitigation plan."),
    ("Χριστίνα Λατρώνη", "0:07:22", "Anna, Stefan — anything else from your side?"),
    ("Anna Schmidt", "0:07:32", "Nothing else. Thank you for the clear update."),
    ("Στέφανος Βάλτερ", "0:07:38", "Same. Will wait for the recap and the formal go/no-go memo."),
    ("Χριστίνα Λατρώνη", "0:07:46", "Thank you everyone. Recap will be in your inbox by end of day."),
]


def main():
    write_docx(
        ROOT / "01_Enerwave_ETL_Status_Sync.docx",
        title="Enerwave - ETL Pipeline Status Sync",
        datetime_line="Wednesday, May 13, 2026, 10:00 AM (EEST)",
        lines=ETL_LINES,
    )
    write_docx(
        ROOT / "02_Enerwave_Fabric_Backend_Review.docx",
        title="Enerwave - Fabric Reporting Backend Review",
        datetime_line="Wednesday, May 13, 2026, 2:00 PM (EEST)",
        lines=FABRIC_LINES,
    )
    write_docx(
        ROOT / "03_LIDL_UAT_Exception_Review.docx",
        title="LIDL - UAT Status & Exception Review",
        datetime_line="Πέμπτη, 14 Μαΐου 2026, 11:00 πμ (EEST)",
        lines=LIDL_LINES,
    )


if __name__ == "__main__":
    main()
